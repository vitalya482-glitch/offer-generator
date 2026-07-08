from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path
import re
import sys
import zipfile
import xml.etree.ElementTree as ET
from typing import Any

try:
    from core.models import OfferContext
except Exception:  # pragma: no cover
    OfferContext = Any  # type: ignore

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None  # type: ignore

try:
    from num2words import num2words
except Exception:  # pragma: no cover
    num2words = None  # type: ignore


BRAND_NAME = "DC Eltek"
PROJECTS_MARKER = "02_Projects"

_XLSX_MAIN_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
_XLSX_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CELL_RE = re.compile(r"^([A-Z]+)(\d+)$")

MONTHS_RU = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}

# Поиск строк выполняется по содержимому листа, без жесткой привязки к номерам строк.
HEADER_ALIASES: dict[str, list[tuple[str, int]]] = {
    "name": [
        ("model", 100),
        ("модель", 100),
        ("наименование", 95),
        ("equipment", 90),
        ("оборудование", 90),
        ("description", 80),
        ("описание", 80),
        ("item", 70),
        ("позиция", 70),
        ("part number", 55),
    ],
    "quantity": [
        ("quantity", 100),
        ("qty", 100),
        ("количество", 100),
        ("кол-во", 100),
        ("кол во", 100),
        ("шт", 75),
        ("pcs", 75),
        ("count", 60),
    ],
    # В калькуляциях может быть несколько строк Price, поэтому итоговую цену за единицу
    # ищем ближе к нижнему итоговому блоку.
    "unit_price": [
        ("total per unit", 120),
        ("итого за единицу", 120),
        ("цена за единицу", 115),
        ("цена за ед", 115),
        ("стоимость за единицу", 110),
        ("стоимость за ед", 110),
        ("unit price", 100),
        ("price per unit", 95),
        ("price", 45),
        ("цена", 45),
    ],
    # Предпочитаем нижнюю итоговую строку TOTAL/ИТОГО.
    "total": [
        ("total", 130),
        ("итого", 130),
        ("общая сумма", 120),
        ("сумма итого", 120),
        ("total with vat", 115),
        ("total per quantity", 90),
        ("сумма", 80),
        ("amount", 70),
        ("стоимость", 60),
    ],
    "vat_percent": [
        ("vat", 100),
        ("vat %", 100),
        ("vat, %", 100),
        ("ндс", 100),
        ("ндс %", 100),
        ("ндс, %", 100),
    ],
}

SERVICE_NAME_TOKENS = {
    "%",
    "total",
    "итого",
    "сумма",
    "quantity",
    "qty",
    "количество",
    "кол-во",
    "price",
    "цена",
    "vat",
    "ндс",
    "ddp",
    "kzt",
    "eur",
    "euro",
    "usd",
    "rate",
    "margin",
}

DEFAULT_TERMS = {
    "payment_terms": "70% предоплата, 30% после уведомления о готовности оборудования к отгрузке.",
    "delivery_time": "Срок поставки уточняется после размещения заказа.",
    "delivery_terms": "DDP Алматы",
    "installation_terms": "Монтажные работы не включены",
    "startup_terms": "Пуско-наладочные работы не включены",
    "offer_validity": "Коммерческое предложение действительно в течение 30 календарных дней.",
}


def extract_client_from_project_path(path_text: str) -> str:
    """Возвращает имя клиента из пути ...\\02_Projects\\КЛИЕНТ\\..."""
    if not path_text:
        return ""

    parts = [part for part in path_text.replace("/", "\\").split("\\") if part]
    for index, part in enumerate(parts):
        if part.lower() == PROJECTS_MARKER.lower() and index + 1 < len(parts):
            return parts[index + 1].strip()

    return ""


def _normalize_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\xa0", " ").strip().lower()
    text = text.replace("ё", "е")
    text = re.sub(r"[\r\n\t]+", " ", text)
    text = re.sub(r"[,:;()\[\]{}]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\xa0", " ").strip()
    return re.sub(r"\s+", " ", text)


def _col_letters_to_index(letters: str) -> int:
    result = 0
    for char in letters:
        result = result * 26 + (ord(char) - ord("A") + 1)
    return result


def _cell_ref_to_row_col(ref: str) -> tuple[int, int]:
    match = _CELL_RE.match(ref.upper())
    if not match:
        raise ValueError(f"Некорректная ссылка Excel-ячейки: {ref}")
    col_letters, row_text = match.groups()
    return int(row_text), _col_letters_to_index(col_letters)


def _to_number(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return float(value)

    text = str(value).strip().replace("\xa0", " ")
    if not text:
        return None

    text = re.sub(r"[^0-9,\.\-]", "", text)
    if not text or text in {"-", ".", ","}:
        return None

    if "," in text and "." in text:
        # 1 234,56 или 1,234.56 — оставляем последний разделитель десятичным.
        if text.rfind(",") > text.rfind("."):
            text = text.replace(".", "").replace(",", ".")
        else:
            text = text.replace(",", "")
    else:
        text = text.replace(",", ".")

    try:
        return float(text)
    except ValueError:
        return None


def format_money(value: float | int | None) -> str:
    """Форматирует денежные суммы для КП: 133 254,00."""
    if value is None:
        value = 0
    amount = round(float(value), 2)
    return f"{amount:,.2f}".replace(",", " ").replace(".", ",")


def format_qty(value: float | int | None) -> str:
    if value is None:
        return "0"
    value = float(value)
    if value.is_integer():
        return str(int(value))
    return f"{value:.2f}".rstrip("0").rstrip(".")


def format_offer_date(dt: datetime | None = None) -> str:
    dt = dt or datetime.now()
    return f"{dt.day} {MONTHS_RU[dt.month]} {dt.year} г."


def sanitize_filename(value: str) -> str:
    bad = '<>:"/\\|?*'
    for ch in bad:
        value = value.replace(ch, "")
    value = "_".join(value.split())
    return value.strip("._ ") or "Client"


def extract_revision_number(value: str) -> int | None:
    text = value or ""
    match = re.search(r"(?:^|[\s_\-\(])(?:v|rev)\s*(\d+)(?:\)|\b|$)", text, re.IGNORECASE)
    if not match:
        return None
    try:
        return int(match.group(1))
    except ValueError:
        return None


def find_next_offer_version(output_dir: Path, client_name: str = "", sheet_name: str = "") -> int:
    output_dir = Path(output_dir)
    max_version = 0
    if output_dir.exists():
        for file_path in output_dir.glob("*.docx"):
            name = file_path.name
            if name.startswith("~$"):
                continue
            match = re.search(r"(?:^|[\s_\-])rev\s*(\d+)(?:\.docx|[\s_\-]|$)", name, re.IGNORECASE)
            if match:
                max_version = max(max_version, int(match.group(1)))
    if max_version > 0:
        return max_version + 1
    return extract_revision_number(sheet_name) or 1


def build_offer_filename(client_name: str, offer_version: int, dt: datetime | None = None) -> str:
    dt = dt or datetime.now()
    client = sanitize_filename(client_name).replace(" ", "_")
    return f"offer_{client}_{dt:%d-%m-%y}_rev{offer_version}.docx"



def currency_name(currency: str) -> str:
    code = (currency or "").upper().strip()
    if code == "KZT":
        return "тенге"
    if code == "EUR":
        return "евро"
    if code == "USD":
        return "долларов США"
    return code.lower() or "валюта не указана"

def _triad_words_ru(number: int, feminine: bool = False) -> str:
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот"]
    tens = ["", "десять", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    ones_m = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    ones_f = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    words: list[str] = []
    h = number // 100
    t = (number // 10) % 10
    o = number % 10
    if h:
        words.append(hundreds[h])
    if t == 1:
        words.append(teens[o])
    else:
        if t:
            words.append(tens[t])
        if o:
            words.append((ones_f if feminine else ones_m)[o])
    return " ".join(words)


def _plural_ru(number: int, one: str, two: str, many: str) -> str:
    n = abs(number) % 100
    n1 = n % 10
    if 11 <= n <= 19:
        return many
    if n1 == 1:
        return one
    if 2 <= n1 <= 4:
        return two
    return many


def number_to_words_ru(number: int) -> str:
    if number == 0:
        return "ноль"
    if number < 0:
        return "минус " + number_to_words_ru(abs(number))

    groups = [
        (10**9, "миллиард", "миллиарда", "миллиардов", False),
        (10**6, "миллион", "миллиона", "миллионов", False),
        (10**3, "тысяча", "тысячи", "тысяч", True),
        (1, "", "", "", False),
    ]
    words: list[str] = []
    remainder = number
    for divisor, one, two, many, feminine in groups:
        value = remainder // divisor
        remainder %= divisor
        if not value:
            continue
        words.append(_triad_words_ru(value, feminine=feminine))
        if divisor != 1:
            words.append(_plural_ru(value, one, two, many))
    return " ".join(part for part in words if part).strip()


def money_in_words(amount: float | int | None, currency: str) -> str:
    whole = int(round(float(amount or 0)))
    cur = (currency or "").upper()
    if cur == "KZT":
        main = "тенге"
        minor = "тиын"
    elif cur == "USD":
        main = "долларов США"
        minor = "центов"
    elif cur == "EUR":
        main = "евро"
        minor = "евроцентов"
    else:
        main = currency_name(currency)
        minor = ""

    if num2words is None:
        words = number_to_words_ru(whole)
    else:
        words = num2words(whole, lang="ru")
    return f"{words} {main} 00 {minor}" if minor else f"{words} {main}"


def _read_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    try:
        raw_xml = archive.read("xl/sharedStrings.xml")
    except KeyError:
        return []

    root = ET.fromstring(raw_xml)
    result: list[str] = []
    for item in root.findall(f"{{{_XLSX_MAIN_NS}}}si"):
        chunks: list[str] = []
        for text_node in item.iter(f"{{{_XLSX_MAIN_NS}}}t"):
            chunks.append(text_node.text or "")
        result.append("".join(chunks))
    return result


def _read_workbook_sheet_paths(archive: zipfile.ZipFile) -> dict[str, str]:
    workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
    rels_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))

    rel_targets: dict[str, str] = {}
    for rel in rels_root.findall(f"{{{_PACKAGE_REL_NS}}}Relationship"):
        rel_id = rel.attrib.get("Id")
        target = rel.attrib.get("Target", "")
        if not rel_id or not target:
            continue
        target = target.lstrip("/")
        if not target.startswith("xl/"):
            target = "xl/" + target
        rel_targets[rel_id] = target

    result: dict[str, str] = {}
    for sheet in workbook_root.findall(f".//{{{_XLSX_MAIN_NS}}}sheet"):
        name = sheet.attrib.get("name", "").strip()
        rel_id = sheet.attrib.get(f"{{{_XLSX_REL_NS}}}id")
        if name and rel_id and rel_id in rel_targets:
            result[name] = rel_targets[rel_id]
    return result



def _read_styles(archive: zipfile.ZipFile) -> dict[int, str]:
    """Возвращает Excel number format по style id, чтобы вытащить валюту из формата ячеек."""
    try:
        root = ET.fromstring(archive.read("xl/styles.xml"))
    except KeyError:
        return {}

    custom_formats: dict[int, str] = {}
    num_fmts = root.find(f"{{{_XLSX_MAIN_NS}}}numFmts")
    if num_fmts is not None:
        for num_fmt in num_fmts.findall(f"{{{_XLSX_MAIN_NS}}}numFmt"):
            try:
                num_fmt_id = int(num_fmt.attrib.get("numFmtId", "0"))
            except ValueError:
                continue
            custom_formats[num_fmt_id] = num_fmt.attrib.get("formatCode", "")

    # Базовые встроенные форматы Excel. Валютные обычно приходят как кастомные,
    # но эти id полезно сохранить для полноты.
    builtin_formats = {
        5: "#,##0_);(#,##0)",
        6: "#,##0_);[Red](#,##0)",
        7: "#,##0.00_);(#,##0.00)",
        8: "#,##0.00_);[Red](#,##0.00)",
        44: "_($* #,##0.00_);_($* (#,##0.00);_($* '-'??_);_(@_)",
    }

    style_formats: dict[int, str] = {}
    cell_xfs = root.find(f"{{{_XLSX_MAIN_NS}}}cellXfs")
    if cell_xfs is None:
        return style_formats

    for index, xf in enumerate(cell_xfs.findall(f"{{{_XLSX_MAIN_NS}}}xf")):
        try:
            num_fmt_id = int(xf.attrib.get("numFmtId", "0"))
        except ValueError:
            continue
        style_formats[index] = custom_formats.get(num_fmt_id, builtin_formats.get(num_fmt_id, ""))
    return style_formats


def _read_sheet_matrix_with_formats(
    calc_path: str | Path,
    sheet_name: str,
) -> tuple[dict[int, dict[int, Any]], dict[int, dict[int, str]], int, int]:
    path = Path(calc_path)
    if path.suffix.lower() not in {".xlsx", ".xlsm"}:
        raise ValueError("DC Eltek пока читает только .xlsx/.xlsm. Старый .xls нужно сохранить как .xlsx.")
    if not path.exists():
        raise FileNotFoundError(f"Excel-файл не найден: {path}")

    with zipfile.ZipFile(path) as archive:
        shared_strings = _read_shared_strings(archive)
        style_formats = _read_styles(archive)
        sheet_paths = _read_workbook_sheet_paths(archive)
        if sheet_name not in sheet_paths:
            available = ", ".join(sheet_paths.keys())
            raise ValueError(f"Лист '{sheet_name}' не найден. Доступные листы: {available}")

        root = ET.fromstring(archive.read(sheet_paths[sheet_name]))

    matrix: dict[int, dict[int, Any]] = {}
    formats: dict[int, dict[int, str]] = {}
    max_row = 0
    max_col = 0

    for cell in root.findall(f".//{{{_XLSX_MAIN_NS}}}c"):
        ref = cell.attrib.get("r", "")
        if not ref:
            continue
        try:
            row, col = _cell_ref_to_row_col(ref)
        except ValueError:
            continue

        style_id_text = cell.attrib.get("s")
        if style_id_text is not None:
            try:
                style_id = int(style_id_text)
                fmt = style_formats.get(style_id, "")
                if fmt:
                    formats.setdefault(row, {})[col] = fmt
            except ValueError:
                pass

        cell_type = cell.attrib.get("t")
        value: Any = None

        if cell_type == "inlineStr":
            chunks = [node.text or "" for node in cell.iter(f"{{{_XLSX_MAIN_NS}}}t")]
            value = "".join(chunks)
        else:
            value_node = cell.find(f"{{{_XLSX_MAIN_NS}}}v")
            if value_node is None or value_node.text is None:
                max_row = max(max_row, row)
                max_col = max(max_col, col)
                continue
            raw_value = value_node.text
            if cell_type == "s":
                index = int(raw_value)
                value = shared_strings[index] if 0 <= index < len(shared_strings) else raw_value
            elif cell_type == "str":
                value = raw_value
            else:
                try:
                    number = float(raw_value)
                    value = int(number) if number.is_integer() else number
                except ValueError:
                    value = raw_value

        if value is not None and value != "":
            matrix.setdefault(row, {})[col] = value
        max_row = max(max_row, row)
        max_col = max(max_col, col)

    return matrix, formats, max_row, max_col


def _read_sheet_matrix(calc_path: str | Path, sheet_name: str) -> tuple[dict[int, dict[int, Any]], int, int]:
    matrix, _formats, max_row, max_col = _read_sheet_matrix_with_formats(calc_path, sheet_name)
    return matrix, max_row, max_col

def _cell_text(matrix: dict[int, dict[int, Any]], row: int, col: int) -> str:
    return _clean_text(matrix.get(row, {}).get(col))


def _cell_number(matrix: dict[int, dict[int, Any]], row: int | None, col: int) -> float | None:
    if row is None:
        return None
    return _to_number(matrix.get(row, {}).get(col))


def _alias_matches(text: str, alias: str) -> bool:
    normalized_alias = _normalize_text(alias)
    if not text or not normalized_alias:
        return False
    if text == normalized_alias:
        return True
    # Для коротких алиасов вроде qty/vat требуем границы слова.
    if len(normalized_alias) <= 4 and re.search(rf"(^|\s){re.escape(normalized_alias)}($|\s|%)", text):
        return True
    return normalized_alias in text


def _detect_best_row(
    matrix: dict[int, dict[int, Any]],
    field: str,
    *,
    min_row: int | None = None,
    max_row: int | None = None,
    prefer_bottom: bool = False,
) -> int | None:
    aliases = HEADER_ALIASES[field]
    candidates: list[tuple[int, int, int, str]] = []

    for row, cols in matrix.items():
        if min_row is not None and row < min_row:
            continue
        if max_row is not None and row > max_row:
            continue

        for col, value in cols.items():
            text = _normalize_text(value)
            if not text:
                continue
            for alias, base_score in aliases:
                if not _alias_matches(text, alias):
                    continue

                score = base_score
                alias_norm = _normalize_text(alias)
                if text == alias_norm:
                    score += 35
                # Заголовки чаще всего в левой части листа.
                score += max(0, 20 - col)
                # Для итоговых строк полезно выбирать нижний блок, для model/quantity — верхний.
                score += row if prefer_bottom else max(0, 200 - row)
                candidates.append((score, row, col, text))

    if not candidates:
        return None

    candidates.sort(reverse=True)
    return candidates[0][1]


def _row_text(matrix: dict[int, dict[int, Any]], row: int) -> str:
    return " ".join(
        _normalize_text(value) for value in matrix.get(row, {}).values()
    ).strip()


def _positive_numeric_count(
    matrix: dict[int, dict[int, Any]],
    row: int,
    start_col: int = 1,
) -> int:
    count = 0
    for col, value in matrix.get(row, {}).items():
        if col < start_col:
            continue
        number = _to_number(value)
        if number is not None and number > 0:
            count += 1
    return count


def _find_exact_label_rows(
    matrix: dict[int, dict[int, Any]],
    labels: tuple[str, ...],
    *,
    min_row: int | None = None,
    max_row: int | None = None,
) -> list[int]:
    normalized = tuple(_normalize_text(label) for label in labels)
    rows: list[int] = []
    for row in sorted(matrix):
        if min_row is not None and row < min_row:
            continue
        if max_row is not None and row > max_row:
            continue
        for value in matrix.get(row, {}).values():
            text = _normalize_text(value)
            if text in normalized:
                rows.append(row)
                break
    return rows


def _find_total_row_after_unit(
    matrix: dict[int, dict[int, Any]],
    unit_price_row: int,
    max_row: int,
) -> int | None:
    """Находит итоговую строку с суммой позиции после нижней цены за единицу."""
    incoterm_rows: list[int] = []
    for row in range(unit_price_row + 1, max_row + 1):
        text = _row_text(matrix, row)
        if not text:
            continue
        if any(_alias_matches(text, alias) for alias in INCOTERM_ALIASES):
            if _positive_numeric_count(matrix, row, start_col=2) >= 1:
                incoterm_rows.append(row)
    if incoterm_rows:
        return incoterm_rows[0]

    total_quantity_rows = _find_exact_label_rows(
        matrix,
        ("total per quantity", "итого за количество", "общая сумма"),
        min_row=unit_price_row + 1,
        max_row=max_row,
    )
    candidates = [
        row
        for row in total_quantity_rows
        if _positive_numeric_count(matrix, row, start_col=2) >= 1
    ]
    return candidates[-1] if candidates else None


def _find_vat_row(
    matrix: dict[int, dict[int, Any]],
    *,
    quantity_row: int,
    unit_price_row: int,
) -> int | None:
    rows: list[int] = []
    for row in range(quantity_row, unit_price_row):
        text = _row_text(matrix, row)
        if not text:
            continue
        if _alias_matches(text, "vat") or _alias_matches(text, "ндс"):
            rows.append(row)
    return rows[-1] if rows else None


def _extract_single_vat_percent(
    matrix: dict[int, dict[int, Any]],
    vat_row: int | None,
) -> float:
    """Берёт один процент НДС из служебной левой части строки VAT/НДС."""
    if not vat_row:
        return 0.0
    for col in sorted(matrix.get(vat_row, {})):
        number = _to_number(matrix[vat_row][col])
        if number is not None and 0 < number <= 100:
            return float(number)
    return 0.0


def detect_offer_layout(calc_path: str | Path, sheet_name: str) -> dict[str, int | None]:
    matrix, _formats, max_row, _max_col = _read_sheet_matrix_with_formats(
        calc_path,
        sheet_name,
    )

    name_row = _detect_best_row(matrix, "name", prefer_bottom=False)
    quantity_row = (
        _detect_best_row(matrix, "quantity", min_row=name_row, prefer_bottom=False)
        if name_row
        else _detect_best_row(matrix, "quantity")
    )

    provisional_total_row = (
        _detect_best_row(matrix, "total", min_row=quantity_row, prefer_bottom=True)
        if quantity_row
        else _detect_best_row(matrix, "total", prefer_bottom=True)
    )

    unit_rows = _find_exact_label_rows(
        matrix,
        (
            "total per unit",
            "итого за единицу",
            "цена за единицу",
            "стоимость за единицу",
        ),
        min_row=quantity_row or name_row or 1,
        max_row=max_row,
    )
    unit_rows = [
        row
        for row in unit_rows
        if _positive_numeric_count(matrix, row, start_col=2) >= 1
    ]

    if unit_rows:
        unit_price_row = unit_rows[-1]
    else:
        unit_max = provisional_total_row - 1 if provisional_total_row else None
        unit_price_row = _detect_best_row(
            matrix,
            "unit_price",
            min_row=quantity_row,
            max_row=unit_max,
            prefer_bottom=True,
        )
        if unit_price_row is None:
            unit_price_row = _detect_best_row(
                matrix,
                "unit_price",
                prefer_bottom=True,
            )

    total_row = provisional_total_row
    if unit_price_row:
        corrected_total_row = _find_total_row_after_unit(
            matrix,
            int(unit_price_row),
            max_row,
        )
        if corrected_total_row:
            total_row = corrected_total_row

    vat_row = None
    if quantity_row and unit_price_row:
        vat_row = _find_vat_row(
            matrix,
            quantity_row=int(quantity_row),
            unit_price_row=int(unit_price_row),
        )
    if vat_row is None and total_row:
        vat_row = _detect_best_row(
            matrix,
            "vat_percent",
            min_row=quantity_row,
            max_row=int(total_row) - 1,
            prefer_bottom=True,
        )
    if vat_row is None:
        vat_row = _detect_best_row(
            matrix,
            "vat_percent",
            prefer_bottom=True,
        )

    return {
        "name_row": name_row,
        "quantity_row": quantity_row,
        "unit_price_row": unit_price_row,
        "total_row": total_row,
        "vat_percent_row": vat_row,
    }

def _is_service_name(name: str) -> bool:
    text = _normalize_text(name)
    if not text:
        return True
    if text in SERVICE_NAME_TOKENS:
        return True
    if text.replace(".", "", 1).isdigit():
        return True
    if len(text) <= 1:
        return True
    return False



def _currency_from_text(text: str) -> str | None:
    text_norm = _normalize_text(text)
    raw = str(text or "").strip().lower()
    if any(term in text_norm for term in ("kzt", "тенге", "тг")) or "₸" in raw:
        return "KZT"
    if any(term in text_norm for term in ("eur", "euro", "евро")) or "€" in raw:
        return "EUR"
    if any(term in text_norm for term in ("usd", "доллар")) or "$" in raw:
        return "USD"
    return None


def _detect_currency(
    matrix: dict[int, dict[int, Any]],
    formats: dict[int, dict[int, str]] | None = None,
    *,
    unit_price_row: int | None = None,
    total_row: int | None = None,
) -> str:
    """Ищет валюту только рядом с ценовыми строками.

    Важно: если валюта не указана в блоке Price per unit / Total per quantity / TOTAL,
    возвращаем пустую строку, а не KZT по умолчанию.
    """
    formats = formats or {}
    candidate_rows: list[int] = []
    for row in (unit_price_row, total_row):
        if row:
            candidate_rows.extend([row - 1, row, row + 1])

    seen: set[int] = set()
    ordered_rows = [row for row in candidate_rows if row and not (row in seen or seen.add(row))]

    for row in ordered_rows:
        row_values = matrix.get(row, {})
        row_formats = formats.get(row, {})
        for col in sorted(set(row_values) | set(row_formats)):
            for source in (row_values.get(col), row_formats.get(col)):
                currency = _currency_from_text(str(source or ""))
                if currency:
                    return currency
    return ""


def detect_dc_eltek_currency(calc_path: str | Path, sheet_name: str) -> str:
    matrix, formats, _max_row, _max_col = _read_sheet_matrix_with_formats(calc_path, sheet_name)
    layout = detect_offer_layout(calc_path, sheet_name)
    return _detect_currency(
        matrix,
        formats,
        unit_price_row=int(layout["unit_price_row"] or 0) if layout.get("unit_price_row") else None,
        total_row=int(layout["total_row"] or 0) if layout.get("total_row") else None,
    )



INCOTERM_ALIASES = (
    "ddp",
    "dap",
    "ddu",
    "exw",
    "fob",
    "cpt",
    "cip",
    "cfr",
    "cif",
    "условия поставки",
    "место поставки",
    "delivery terms",
    "delivery condition",
)

INSTALLATION_ALIASES = (
    "installation",
    "install",
    "монтаж",
    "монтажные работы",
)

STARTUP_ALIASES = (
    "start-up",
    "startup",
    "commissioning",
    "commission",
    "пнр",
    "пуско",
    "налад",
)

INSPECTION_ALIASES = (
    "inspection",
    "site survey",
    "survey",
    "обследование",
    "инспекция",
    "инспекция объекта",
    "выезд на объект",
    "аудит объекта",
)

SPECIAL_TERM_ALIASES = (
    "special term",
    "special terms",
    "special condition",
    "special conditions",
    "финансирование",
    "финанс",
    "особые условия",
    "особое условие",
)


def _row_label(matrix: dict[int, dict[int, Any]], row: int) -> str:
    cols = matrix.get(row, {})
    if not cols:
        return ""
    # Обычно названия расчетных строк стоят в левой части листа.
    for col in sorted(cols):
        text = _clean_text(cols[col])
        if text:
            return text
    return ""


def _row_contains_any(matrix: dict[int, dict[int, Any]], row: int, aliases: tuple[str, ...]) -> bool:
    row_text = " ".join(_normalize_text(value) for value in matrix.get(row, {}).values())
    return any(_alias_matches(row_text, alias) for alias in aliases)


def _find_rows_by_aliases(matrix: dict[int, dict[int, Any]], aliases: tuple[str, ...]) -> list[int]:
    rows: list[int] = []
    for row in sorted(matrix):
        if _row_contains_any(matrix, row, aliases):
            rows.append(row)
    return rows


def _sum_rows_for_item_columns(
    matrix: dict[int, dict[int, Any]],
    rows: list[int],
    item_columns: list[int],
) -> float:
    total = 0.0
    for row in rows:
        for col in item_columns:
            value = _cell_number(matrix, row, col)
            if value is not None and value > 0:
                total += float(value)
    return total


def _detect_delivery_terms(
    matrix: dict[int, dict[int, Any]],
    *,
    total_row: int | None,
) -> str:
    candidates: list[tuple[int, int, str]] = []
    for row, cols in matrix.items():
        if total_row is not None and row > total_row:
            continue
        for col, value in cols.items():
            text = _clean_text(value)
            normalized = _normalize_text(text)
            if not normalized:
                continue
            if any(_alias_matches(normalized, alias) for alias in INCOTERM_ALIASES):
                # Предпочитаем ближайшую к итоговому блоку строку, затем левую часть листа.
                row_score = row if total_row is None else (1000 - abs(total_row - row))
                col_score = max(0, 50 - col)
                candidates.append((row_score + col_score, row, text))

    if not candidates:
        return ""
    candidates.sort(reverse=True)
    return candidates[0][2]


def _status_from_cost(total: float, *, included_word: str, not_included_word: str, currency: str) -> str:
    if total > 0.01:
        return f"{included_word}, сумма {format_money(total)} {currency}"
    return not_included_word


def _percent_for_row_column(matrix: dict[int, dict[int, Any]], row: int, item_col: int) -> float | None:
    # В расчетах Eltek проценты часто стоят в соседней колонке слева от суммы позиции:
    # B=%, C=сумма; D=%, E=сумма и т.д. Но оставляем запасные варианты.
    for col in (item_col - 1, item_col + 1, item_col):
        if col <= 0:
            continue
        value = _cell_number(matrix, row, col)
        if value is not None and 0 < value <= 100:
            return float(value)
    return None


def _detect_special_terms(
    matrix: dict[int, dict[int, Any]],
    item_columns: list[int],
    currency: str,
) -> dict[str, Any]:
    rows = _find_rows_by_aliases(matrix, SPECIAL_TERM_ALIASES)
    total = _sum_rows_for_item_columns(matrix, rows, item_columns)

    rates: list[float] = []
    for row in rows:
        for col in item_columns:
            rate = _percent_for_row_column(matrix, row, col)
            if rate is not None:
                rates.append(rate)

    unique_rates = sorted({round(rate, 4) for rate in rates})
    rate_label = ", ".join(f"{rate:g}%" for rate in unique_rates)

    if total > 0.01 or unique_rates:
        if rate_label:
            text = f"{rate_label}, сумма {format_money(total)} {currency}" if total > 0.01 else rate_label
        else:
            text = f"сумма {format_money(total)} {currency}"
    else:
        text = "не заложено"

    return {
        "rows": rows,
        "total": total,
        "rates": unique_rates,
        "rate_label": rate_label,
        "text": text,
    }


def detect_additional_terms(
    matrix: dict[int, dict[int, Any]],
    layout: dict[str, int | None],
    items: list[dict[str, Any]],
    currency: str,
) -> dict[str, Any]:
    item_columns = [int(item.get("source_column") or 0) for item in items if int(item.get("source_column") or 0) > 0]
    total_row = int(layout.get("total_row") or 0) or None

    installation_rows = _find_rows_by_aliases(matrix, INSTALLATION_ALIASES)
    startup_rows = _find_rows_by_aliases(matrix, STARTUP_ALIASES)
    combined_rows = sorted(set(installation_rows + startup_rows))
    inspection_rows = _find_rows_by_aliases(matrix, INSPECTION_ALIASES)

    installation_startup_total = _sum_rows_for_item_columns(matrix, combined_rows, item_columns)
    inspection_total = _sum_rows_for_item_columns(matrix, inspection_rows, item_columns)
    special_terms = _detect_special_terms(matrix, item_columns, currency)

    delivery_terms = _detect_delivery_terms(matrix, total_row=total_row)

    installation_pnr_status = _status_from_cost(
        installation_startup_total,
        included_word="включены",
        not_included_word="не включены",
        currency=currency,
    )
    inspection_status = _status_from_cost(
        inspection_total,
        included_word="заложена",
        not_included_word="не заложена",
        currency=currency,
    )

    return {
        "delivery_terms": delivery_terms,
        "currency": currency,
        "installation_rows": combined_rows,
        "installation_startup_total": installation_startup_total,
        "installation_pnr_status": installation_pnr_status,
        "installation_terms": "Монтажные работы включены" if installation_startup_total > 0.01 else "Монтажные работы не включены",
        "startup_terms": "Пуско-наладочные работы включены" if installation_startup_total > 0.01 else "Пуско-наладочные работы не включены",
        "inspection_rows": inspection_rows,
        "inspection_total": inspection_total,
        "inspection_status": inspection_status,
        "special_terms": special_terms,
    }


def read_dc_eltek_offer_items(
    calc_path: str | Path,
    sheet_name: str,
    currency_override: str | None = None,
) -> dict[str, Any]:
    """Читает выбранный лист DC Eltek и возвращает позиции + итоги.

    Парсер не привязан к конкретным строкам/колонкам. Он ищет строки по заголовкам
    Model/Quantity/Total per unit/TOTAL/VAT и затем проходит все колонки слева направо.
    Валюта должна быть найдена в ценовом блоке или явно выбрана пользователем.
    """
    matrix, formats, _max_row, max_col = _read_sheet_matrix_with_formats(calc_path, sheet_name)
    layout = detect_offer_layout(calc_path, sheet_name)

    required = ["name_row", "quantity_row", "unit_price_row", "total_row"]
    missing = [key for key in required if not layout.get(key)]
    if missing:
        raise ValueError(
            "Не удалось определить структуру листа DC Eltek. Не найдены строки: "
            + ", ".join(missing)
        )

    name_row = int(layout["name_row"] or 0)
    quantity_row = int(layout["quantity_row"] or 0)
    unit_price_row = int(layout["unit_price_row"] or 0)
    total_row = int(layout["total_row"] or 0)
    vat_percent_row = int(layout.get("vat_percent_row") or 0) or None
    common_vat_percent = _extract_single_vat_percent(matrix, vat_percent_row)
    detected_currency = _detect_currency(
        matrix,
        formats,
        unit_price_row=unit_price_row,
        total_row=total_row,
    )
    currency = (currency_override or detected_currency or "").upper().strip()

    items: list[dict[str, Any]] = []

    for col in range(1, max_col + 1):
        name = _cell_text(matrix, name_row, col)
        if _is_service_name(name):
            continue

        quantity = _cell_number(matrix, quantity_row, col)
        unit_price = _cell_number(matrix, unit_price_row, col)
        total = _cell_number(matrix, total_row, col)
        vat_percent = common_vat_percent

        if quantity is None or quantity <= 0:
            continue
        if (unit_price is None or unit_price <= 0) and (total is None or total <= 0):
            continue

        if (total is None or total <= 0) and unit_price is not None:
            total = unit_price * quantity
        if (unit_price is None or unit_price <= 0) and total is not None and quantity:
            unit_price = total / quantity

        vat_percent = vat_percent or 0.0
        total_with_vat = float(total or 0.0)
        if vat_percent > 0:
            total_without_vat = total_with_vat / (1 + vat_percent / 100)
            vat_amount = total_with_vat - total_without_vat
        else:
            total_without_vat = total_with_vat
            vat_amount = 0.0

        items.append(
            {
                "number": len(items) + 1,
                "name": name,
                "quantity": quantity,
                "unit_price": float(unit_price or 0.0),
                "total_without_vat": total_without_vat,
                "vat_percent": vat_percent,
                "vat_amount": vat_amount,
                "total": total_with_vat,
                "currency": currency,
                "detected_currency": detected_currency,
                "source_column": col,
            }
        )

    summary = summarize_items(items, currency=currency)
    summary["detected_currency"] = detected_currency
    summary["currency_source"] = "manual" if currency_override else ("auto" if detected_currency else "missing")
    meta = detect_additional_terms(matrix, layout, items, currency)
    return {
        "layout": layout,
        "items": items,
        "summary": summary,
        "meta": meta,
    }

def summarize_items(items: list[dict[str, Any]], currency: str = "KZT") -> dict[str, Any]:
    quantity_total = sum(float(item.get("quantity") or 0) for item in items)
    total_without_vat = sum(float(item.get("total_without_vat") or 0) for item in items)
    vat_amount = sum(float(item.get("vat_amount") or 0) for item in items)
    total = sum(float(item.get("total") or 0) for item in items)

    vat_rates = sorted({round(float(item.get("vat_percent") or 0), 4) for item in items})
    vat_label = ", ".join(f"{rate:g}%" for rate in vat_rates) if vat_rates else "0%"

    return {
        "positions_count": len(items),
        "quantity_total": quantity_total,
        "total_without_vat": total_without_vat,
        "vat_amount": vat_amount,
        "total": total,
        "vat_rates": vat_rates,
        "vat_label": vat_label,
        "currency": currency,
    }


def _iter_candidate_roots() -> list[Path]:
    roots: list[Path] = []
    for candidate in (
        Path.cwd(),
        Path(sys.executable).resolve().parent if getattr(sys, "frozen", False) else None,
        Path(__file__).resolve().parents[1] if "__file__" in globals() else None,
        Path(__file__).resolve().parents[2] if "__file__" in globals() and len(Path(__file__).resolve().parents) > 2 else None,
    ):
        if candidate and candidate not in roots:
            roots.append(candidate)
    return roots


def find_default_dc_eltek_template() -> str:
    """Ищет встроенный/локальный шаблон DC Eltek, если пользователь не выбрал файл вручную."""
    names = (
        "Offer_Company_02-07-26_TAGGED_Eltek.docx",
        "dc_eltek_template.docx",
        "dc_eltek_offer_template.docx",
    )
    checked_dirs: list[Path] = []
    for root in _iter_candidate_roots():
        for folder in (root / "templates" / "dc_eltek", root / "templates"):
            if folder in checked_dirs:
                continue
            checked_dirs.append(folder)
            for name in names:
                candidate = folder / name
                if candidate.exists():
                    return str(candidate)
            if folder.exists():
                for candidate in folder.glob("*.docx"):
                    if "eltek" in candidate.name.lower() and not candidate.name.startswith("~$"):
                        return str(candidate)
    return ""


def _context_get(context: OfferContext | dict[str, Any], key: str, default: Any = "") -> Any:
    if isinstance(context, dict):
        return context.get(key, default)
    return getattr(context, key, default)


def _context_values(context: OfferContext | dict[str, Any]) -> dict[str, Any]:
    project_dir = str(_context_get(context, "project_dir", "") or "")
    client = str(_context_get(context, "client", "") or _context_get(context, "client_name", "") or "")
    if not client:
        client = extract_client_from_project_path(project_dir)
    calc_path = str(_context_get(context, "calc_path", "") or "")
    sheet_name = str(_context_get(context, "sheet_name", "") or "")
    template_path = str(_context_get(context, "template_path", "") or "") or find_default_dc_eltek_template()
    output_dir = str(_context_get(context, "output_dir", "") or project_dir or Path(calc_path).parent)

    signer_name = str(_context_get(context, "signer_name", "") or "Сания Санаткызы")
    signer_position = str(_context_get(context, "signer_position", "") or "Коммерческий директор")
    manager_name = str(_context_get(context, "manager_name", "") or "")
    manager_position = str(_context_get(context, "manager_position", "") or "")
    manager_email = str(_context_get(context, "manager_email", "") or "")
    manager_phone = str(_context_get(context, "manager_phone", "") or "")

    return {
        "project_dir": project_dir,
        "client": client,
        "calc_path": calc_path,
        "sheet_name": sheet_name,
        "template_path": template_path,
        "output_dir": output_dir,
        "signer_name": signer_name,
        "signer_position": signer_position,
        "manager_name": manager_name,
        "manager_position": manager_position,
        "manager_email": manager_email,
        "manager_phone": manager_phone,
        "currency": str(_context_get(context, "currency", "") or "").upper().strip(),
    }


def build_intro_text(summary: dict[str, Any]) -> str:
    qty = format_qty(summary.get("quantity_total"))
    positions_count = int(summary.get("positions_count") or 0)
    return (
        "В ответ на Ваш запрос направляем коммерческое предложение на поставку "
        f"оборудования систем питания постоянного тока Eltek: {positions_count} поз. "
        f"общим количеством {qty} шт."
    )


def build_total_price_block(summary: dict[str, Any]) -> str:
    currency = str(summary.get("currency") or "")
    total = float(summary.get("total") or 0.0)
    vat_amount = float(summary.get("vat_amount") or 0.0)
    vat_label = str(summary.get("vat_label") or "0%")
    vat_text = "без учета НДС" if abs(vat_amount) < 0.01 else f"с учетом НДС {vat_label}"
    return (
        f"{format_money(total)} {currency_name(currency)} "
        f"({money_in_words(total, currency)}), {vat_text}."
    )


def build_currency_terms(summary: dict[str, Any]) -> str:
    currency = str(summary.get("currency") or "").upper()
    if currency == "KZT":
        return "Стоимость указана в тенге."
    if currency == "EUR":
        return "Взаиморасчет осуществляется в тенге по курсу АО Банк ЦентрКредит на день оплаты."
    if currency == "USD":
        return "Взаиморасчет осуществляется в тенге по курсу банка на день оплаты."
    return "Стоимость указана в валюте коммерческого предложения."


def build_offer_items(parsed: dict[str, Any]) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    for item in parsed.get("items", []):
        result.append(
            {
                "item_no": item.get("number", ""),
                "item_name": item.get("name", ""),
                "item_qty": format_qty(item.get("quantity")),
                "item_unit_price": format_money(item.get("unit_price")),
                "item_total": format_money(item.get("total")),
            }
        )
    return result


def build_replacements(context_values: dict[str, Any], parsed: dict[str, Any], offer_version: int | None = None) -> dict[str, Any]:
    summary = parsed.get("summary", {})
    meta = parsed.get("meta", {})
    special_terms = meta.get("special_terms", {}) if isinstance(meta.get("special_terms", {}), dict) else {}

    currency = str(summary.get("currency") or meta.get("currency") or "")
    cur_name = currency_name(currency)
    terms = DEFAULT_TERMS

    delivery_terms = str(meta.get("delivery_terms") or terms["delivery_terms"])
    installation_terms = str(meta.get("installation_terms") or terms["installation_terms"])
    startup_terms = str(meta.get("startup_terms") or terms["startup_terms"])
    inspection_terms = str(meta.get("inspection_status") or "не заложена")
    financing_terms = str(special_terms.get("text") or "не заложено")

    return {
        "{{offer_date}}": format_offer_date(),
        "{{offer_version}}": str(offer_version or 1),
        "{{client_company_full}}": context_values.get("client", ""),
        "{{intro_text}}": build_intro_text(summary),
        "{{unit_price_header}}": f"Цена за единицу, {cur_name}",
        "{{total_price_header}}": f"Сумма, {cur_name}",
        "{{total_label}}": "ИТОГО",
        "{{grand_total}}": format_money(summary.get("total")),
        "{{total_price_block}}": build_total_price_block(summary),
        "{{payment_terms}}": terms["payment_terms"],
        "{{delivery_time}}": terms["delivery_time"],
        "{{delivery_terms}}": delivery_terms,
        "{{installation_terms}}": installation_terms,
        "{{startup_terms}}": startup_terms,
        "{{offer_validity}}": terms["offer_validity"],
        "{{currency_terms}}": build_currency_terms(summary),
        "{{signer_name}}": context_values.get("signer_name", ""),
        "{{signer_position}}": context_values.get("signer_position", ""),
        "{{manager_name}}": context_values.get("manager_name", ""),
        "{{manager_position}}": context_values.get("manager_position", ""),
        "{{manager_email}}": context_values.get("manager_email", ""),
        "{{manager_phone}}": context_values.get("manager_phone", ""),

        # Дополнительные теги на будущие версии шаблонов. Если тегов нет в docx — они просто игнорируются.
        "{{delivery_place_terms}}": delivery_terms,
        "{{currency}}": currency,
        "{{currency_code}}": currency,
        "{{mounting_pnr_status}}": str(meta.get("installation_pnr_status") or "не включены"),
        "{{installation_pnr_status}}": str(meta.get("installation_pnr_status") or "не включены"),
        "{{inspection_terms}}": inspection_terms,
        "{{inspection_status}}": inspection_terms,
        "{{special_terms}}": financing_terms,
        "{{financing_terms}}": financing_terms,
        "{{special_financing_terms}}": financing_terms,
    }

def _docx_to_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def _paragraph_full_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def _replace_tag_across_runs(paragraph, tag: str, value: Any) -> bool:
    if not paragraph.runs or not tag:
        return False
    full_text = _paragraph_full_text(paragraph)
    start = full_text.find(tag)
    if start < 0:
        return False
    end = start + len(tag)
    spans: list[tuple[int, int, Any]] = []
    pos = 0
    for run in paragraph.runs:
        run_start = pos
        run_end = pos + len(run.text)
        spans.append((run_start, run_end, run))
        pos = run_end
    involved = [item for item in spans if item[1] > start and item[0] < end]
    if not involved:
        return False
    first_start, _first_end, first_run = involved[0]
    _last_start, _last_end, last_run = involved[-1]
    prefix = first_run.text[: max(0, start - first_start)]
    suffix = last_run.text[max(0, end - _last_start) :]
    replacement = _docx_to_text(value)
    if first_run is last_run:
        first_run.text = prefix + replacement + suffix
    else:
        first_run.text = prefix + replacement
        for _, _, middle_run in involved[1:-1]:
            middle_run.text = ""
        last_run.text = suffix
    return True


def _replace_in_paragraph(paragraph, replacements: dict[str, Any]) -> None:
    if not paragraph.runs:
        return
    changed = True
    while changed:
        changed = False
        full_text = _paragraph_full_text(paragraph)
        for key, value in replacements.items():
            if key and key in full_text:
                changed = _replace_tag_across_runs(paragraph, key, value) or changed
                full_text = _paragraph_full_text(paragraph)


def _replace_tags(doc, replacements: dict[str, Any]) -> None:
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)


def _set_cell_keep_style(cell, value: Any) -> None:
    value = _docx_to_text(value)
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)
    for extra_paragraph in cell.paragraphs[1:]:
        for run in extra_paragraph.runs:
            run.text = ""


def _row_contains_tags(row, tags: list[str]) -> bool:
    row_text = "\n".join(cell.text for cell in row.cells)
    return any(tag in row_text for tag in tags)


def _clone_row_after(table, source_row, after_row):
    new_tr = deepcopy(source_row._tr)
    after_row._tr.addnext(new_tr)
    return table.rows[after_row._index + 1]


def _remove_row(table, row) -> None:
    table._tbl.remove(row._tr)


def _fill_equipment_table(doc, items: list[dict[str, Any]], total_label: str, grand_total: str) -> None:
    item_tags = ["{{item_no}}", "{{item_name}}", "{{item_qty}}", "{{item_unit_price}}", "{{item_total}}"]
    for table in doc.tables:
        template_row = None
        total_row = None
        for row in table.rows:
            if _row_contains_tags(row, item_tags):
                template_row = row
            if _row_contains_tags(row, ["{{total_label}}", "{{grand_total}}"]):
                total_row = row
        if template_row is None:
            continue

        insert_after = template_row
        for item in items:
            new_row = _clone_row_after(table, template_row, insert_after)
            insert_after = new_row
            values = {
                "{{item_no}}": item.get("item_no", ""),
                "{{item_name}}": item.get("item_name", ""),
                "{{item_qty}}": item.get("item_qty", ""),
                "{{item_unit_price}}": item.get("item_unit_price", ""),
                "{{item_total}}": item.get("item_total", ""),
            }
            for cell in new_row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, values)
                for tag, replacement_value in values.items():
                    if tag in cell.text:
                        _set_cell_keep_style(cell, replacement_value)
                        break

        _remove_row(table, template_row)

        if total_row is not None:
            values = {"{{total_label}}": total_label, "{{grand_total}}": grand_total}
            for cell in total_row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, values)
                for tag, replacement_value in values.items():
                    if tag in cell.text:
                        _set_cell_keep_style(cell, replacement_value)
                        break
        return


def _render_dc_eltek_docx(
    template_path: str | Path,
    output_path: str | Path,
    replacements: dict[str, Any],
    items: list[dict[str, Any]],
) -> Path:
    if Document is None:
        raise RuntimeError("Для формирования КП нужен пакет python-docx.")

    template_path = Path(template_path)
    output_path = Path(output_path)
    if not template_path.exists():
        raise FileNotFoundError(f"Шаблон КП не найден: {template_path}")

    doc = Document(str(template_path))
    _fill_equipment_table(
        doc=doc,
        items=items,
        total_label=str(replacements.get("{{total_label}}", "ИТОГО")),
        grand_total=str(replacements.get("{{grand_total}}", "")),
    )
    _replace_tags(doc, replacements)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path = _get_unique_output_path(output_path)
    doc.save(str(output_path))
    return output_path


def _get_unique_output_path(path: Path) -> Path:
    if not path.exists():
        return path
    parent = path.parent
    stem = path.stem
    suffix = path.suffix
    counter = 1
    while True:
        candidate = parent / f"{stem}_{counter}{suffix}"
        if not candidate.exists():
            return candidate
        counter += 1



def make_offer(context: OfferContext | dict[str, Any]) -> Path:
    values = _context_values(context)
    calc_path = values["calc_path"]
    sheet_name = values["sheet_name"]
    template_path = values["template_path"]
    output_dir = Path(values["output_dir"] or Path(calc_path).parent)
    client = values["client"] or "DC_Eltek"
    currency_override = values.get("currency", "")

    if not calc_path:
        raise ValueError("Не выбран Excel calc.")
    if not sheet_name:
        raise ValueError("Не выбран лист для КП.")
    if not template_path:
        raise ValueError("Не выбран шаблон КП. Выберите .docx вручную или положите его в templates/dc_eltek.")

    parsed = read_dc_eltek_offer_items(calc_path, sheet_name, currency_override=currency_override)
    if not parsed.get("items"):
        raise ValueError("В выбранном листе Excel не найдены позиции для КП.")

    currency = str(parsed.get("summary", {}).get("currency") or "").upper().strip()
    if not currency:
        raise ValueError(
            "Валюта не указана. Укажите валюту в расчёте рядом со строками "
            "Price per unit / Total per quantity / TOTAL или выберите валюту на вкладке DC Eltek."
        )

    offer_version = find_next_offer_version(output_dir, client, sheet_name)
    replacements = build_replacements(values, parsed, offer_version=offer_version)
    items = build_offer_items(parsed)
    filename = build_offer_filename(client, offer_version)
    output_path = output_dir / filename
    return _render_dc_eltek_docx(template_path, output_path, replacements, items)



def preview(context: OfferContext | dict[str, Any]) -> str:
    values = _context_values(context)
    calc_path = values["calc_path"]
    sheet_name = values["sheet_name"]
    currency_override = values.get("currency", "")

    if not calc_path or not sheet_name:
        return "Выберите Excel calc и лист для КП — после этого здесь появятся позиции, валюта, условия и итоги."

    try:
        parsed = read_dc_eltek_offer_items(calc_path, sheet_name, currency_override=currency_override)
    except Exception as exc:
        return f"Не удалось прочитать расчёт: {exc}"

    items = parsed["items"]
    summary = parsed["summary"]
    meta = parsed.get("meta", {})
    special_terms = meta.get("special_terms", {}) if isinstance(meta.get("special_terms", {}), dict) else {}
    currency = str(summary.get("currency") or "").upper().strip()
    detected_currency = str(summary.get("detected_currency") or "").upper().strip()

    lines: list[str] = ["Позиции:"]
    for item in items:
        lines.append(
            f"{item['number']}. {item['name']} — {format_qty(item['quantity'])} шт × "
            f"{format_money(item['unit_price'])} = {format_money(item['total'])}"
        )

    lines.extend(
        [
            "",
            "Итоги:",
            f"Количество позиций: {summary['positions_count']}",
            f"Общее количество, шт: {format_qty(summary['quantity_total'])}",
            f"Сумма без НДС: {format_money(summary['total_without_vat'])}",
            f"НДС: {format_money(summary['vat_amount'])} ({summary['vat_label']})",
            f"Общая сумма с НДС: {format_money(summary['total'])}",
            f"Валюта: {currency or 'не указана'}" + (f" (авто: {detected_currency})" if detected_currency and currency_override else ""),
            "",
            "Дополнительно:",
            f"Место и условия поставки: {meta.get('delivery_terms') or '-'}",
            f"Монтаж и ПНР: {meta.get('installation_pnr_status') or 'не включены'}",
            f"Инспекция объекта: {meta.get('inspection_status') or 'не заложена'}",
            f"Особые условия / финансирование: {special_terms.get('text') or 'не заложено'}",
        ]
    )

    if not currency:
        lines.extend(
            [
                "",
                "ВНИМАНИЕ: валюта не указана в ценовом блоке расчёта.",
                "Перед формированием КП выберите валюту вручную на вкладке DC Eltek.",
            ]
        )

    return "\n".join(lines)

