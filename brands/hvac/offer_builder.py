from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Mapping, Sequence

from docx import Document

from .excel_reader import HVACPosition, format_money, format_qty

TAG_RE = re.compile(r"\{\{\s*([A-Z0-9_]+)\s*\}\}")


def build_hvac_offer(
    template_path: str | Path,
    output_path: str | Path,
    fields: Mapping[str, Any],
    variant1_items: Sequence[HVACPosition | Mapping[str, Any]],
    variant2_items: Sequence[HVACPosition | Mapping[str, Any]],
    max_items: int = 8,
) -> Path:
    """Render HVAC DOCX offer by replacing {{TAGS}} in template.

    The template is intentionally simple and fixed-width. Empty rows are blanked.
    No specification files are attached.
    """
    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"Шаблон КП не найден: {template}")

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    tags = make_hvac_tags(fields, variant1_items, variant2_items, max_items=max_items)
    doc = Document(str(template))
    _replace_tags_in_document(doc, tags)
    doc.save(str(output))
    return output


def make_hvac_tags(
    fields: Mapping[str, Any],
    variant1_items: Sequence[HVACPosition | Mapping[str, Any]],
    variant2_items: Sequence[HVACPosition | Mapping[str, Any]],
    max_items: int = 8,
) -> dict[str, str]:
    tags: dict[str, str] = {str(k).upper(): _text(v) for k, v in fields.items()}

    tags.setdefault("CITY", "Алматы")
    tags.setdefault("OFFER_VERSION", "1")
    tags.setdefault("CURRENCY", "ЕВРО")
    tags.setdefault("CURRENCY_RATE_TEXT", "Взаиморасчет осуществляется в тенге по курсу АО Банк ЦентрКредит на день оплаты.")
    tags.setdefault("PAYMENT_TERMS", "70% предоплата, 30% после поставки")
    tags.setdefault("VALIDITY_TEXT", "Предложение действительно в течение 30 дней.")
    tags.setdefault("V1_NOTE_1", "*Инжиниринг включен.")
    tags.setdefault("V1_NOTE_2", "*Монтажные и пуско-наладочные работы не включены.")
    tags.setdefault("V2_NOTE_1", "*Инжиниринг включен.")
    tags.setdefault("V2_NOTE_2", "*Монтажные и пуско-наладочные работы не включены.")

    _fill_basis_tags(tags, fields)
    _fill_deviation_tags(tags, fields)
    _fill_variant_tags(tags, "V1", variant1_items, max_items)
    _fill_variant_tags(tags, "V2", variant2_items, max_items)
    return tags


def _fill_basis_tags(tags: dict[str, str], fields: Mapping[str, Any]) -> None:
    basis_docs = fields.get("BASIS_DOCS") or fields.get("basis_docs") or []
    if isinstance(basis_docs, str):
        basis_docs = [x.strip() for x in basis_docs.splitlines() if x.strip()]
    for i in range(1, 7):
        tags.setdefault(f"BASIS_DOC_{i}", _text(basis_docs[i - 1]) if i <= len(basis_docs) else "")


def _fill_deviation_tags(tags: dict[str, str], fields: Mapping[str, Any]) -> None:
    tags.setdefault("DEV_1_NO", "1")
    tags.setdefault(
        "DEV_1_NAME",
        "Датчики газа/дыма;\nВ стоимость не включены кабели, медные трубы, распредшкаф для питания, опоры, трубы",
    )
    tags.setdefault("DEV_1_COMMENT", "не входит в объем поставки оборудования ОВКВ")


def _fill_variant_tags(
    tags: dict[str, str],
    prefix: str,
    items: Sequence[HVACPosition | Mapping[str, Any]],
    max_items: int,
) -> None:
    total = 0.0
    for idx in range(1, max_items + 1):
        if idx <= len(items):
            item = _as_item(items[idx - 1])
            amount = item.get("amount")
            try:
                total += float(amount or 0)
            except (TypeError, ValueError):
                pass
            tags[f"{prefix}_ITEM_{idx}_NO"] = str(idx)
            tags[f"{prefix}_ITEM_{idx}_NAME"] = _text(item.get("name"))
            tags[f"{prefix}_ITEM_{idx}_QTY"] = format_qty(item.get("qty"))
            tags[f"{prefix}_ITEM_{idx}_AMOUNT"] = format_money(amount)
        else:
            tags[f"{prefix}_ITEM_{idx}_NO"] = ""
            tags[f"{prefix}_ITEM_{idx}_NAME"] = ""
            tags[f"{prefix}_ITEM_{idx}_QTY"] = ""
            tags[f"{prefix}_ITEM_{idx}_AMOUNT"] = ""

    tags.setdefault(f"{prefix}_TOTAL_LABEL", "Стоимость без учёта НДС, EUR")
    tags[f"{prefix}_TOTAL_AMOUNT"] = format_money(total)


def _as_item(item: HVACPosition | Mapping[str, Any]) -> dict[str, Any]:
    if isinstance(item, HVACPosition):
        return {"name": item.name, "qty": item.qty, "amount": item.amount}
    return dict(item)


def _replace_tags_in_document(doc: Document, tags: Mapping[str, str]) -> None:
    for paragraph in doc.paragraphs:
        _replace_tags_in_paragraph(paragraph, tags)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_tags_in_paragraph(paragraph, tags)

    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in header_footer.paragraphs:
                _replace_tags_in_paragraph(paragraph, tags)
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _replace_tags_in_paragraph(paragraph, tags)


def _replace_tags_in_paragraph(paragraph, tags: Mapping[str, str]) -> None:
    original = paragraph.text
    if "{{" not in original:
        return

    def repl(match: re.Match[str]) -> str:
        key = match.group(1).strip().upper()
        return _text(tags.get(key, ""))

    replaced = TAG_RE.sub(repl, original)
    if replaced == original:
        return

    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replaced)


def _text(value: Any) -> str:
    if value is None:
        return ""
    return str(value)
