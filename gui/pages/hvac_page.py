from __future__ import annotations

from collections import Counter
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable
import os
import re

from docx import Document
from PySide6.QtCore import Qt
from PySide6.QtWidgets import (
    QAbstractItemView,
    QCheckBox,
    QComboBox,
    QFileDialog,
    QGridLayout,
    QHBoxLayout,
    QHeaderView,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPushButton,
    QTableWidget,
    QTableWidgetItem,
    QVBoxLayout,
    QWidget,
)

from brands.hvac.template_finder import find_default_hvac_template
from core.excel_calc_parser import (
    CalcItem,
    CalcParseResult,
    format_money,
    format_qty,
    parse_calculation,
    read_sheet_names,
)
from core.project_scanner import clear_scan_cache, scan_project_files
from gui.path_helpers import extract_client_from_project_dir, infer_output_dir


_TAG_RE = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}")
_ITEM_TAGS = {"item_no", "item_name", "item_qty", "item_unit_price", "item_total"}
_FROM_CLIENT_ALIASES = (
    "from client",
    "from customer",
    "client data",
    "customer data",
    "от клиента",
    "от заказчика",
    "исходные данные",
    "входящие от клиента",
)
_IGNORED_BASIS_FILES = {"thumbs.db", "desktop.ini", ".ds_store"}


class HVACPage(QWidget):
    """HVAC: calculation Excel -> tagged DOCX offer.

    The page only contains HVAC-specific inputs. Common data is taken from the
    main window:
    - signer from the selected sidebar radio button;
    - manager from Settings;
    - date is generated automatically;
    - client is extracted from the selected project path;
    - basis documents are collected from the project's From Client folder.
    """

    brand_name = "HVAC"

    def __init__(self, owner) -> None:
        super().__init__(owner)
        self.owner = owner
        self.settings = owner.settings
        self._updating_path_display = False

        self.project_dir_path = self._saved("hvac/project_dir", self._saved("project_dir", ""))
        self.output_dir_path = self._saved("hvac/output_dir", "")
        self.client_company = ""
        self.from_client_dir: Path | None = None
        self.basis_files: list[str] = []
        self.parse_result: CalcParseResult | None = None
        self.items: list[CalcItem] = []
        self.item_checks: list[QCheckBox] = []

        self._build_ui()
        self._load_saved_values()
        self._load_default_template()

        if self.project_dir_path and Path(self.project_dir_path).is_dir():
            self.scan_project(force=False)
        else:
            self._refresh_auto_data()
            self._update_status()

    # ------------------------------------------------------------------ UI
    def _build_ui(self) -> None:
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(14)

        header = QHBoxLayout()
        header_text = QVBoxLayout()
        title = QLabel("HVAC: новое коммерческое предложение")
        title.setObjectName("PageTitle")
        subtitle = QLabel(
            "Выберите папку проекта. Заказчик, основания, подписант и исполнитель "
            "заполняются автоматически."
        )
        subtitle.setObjectName("PageSubtitle")
        subtitle.setWordWrap(True)
        header_text.addWidget(title)
        header_text.addWidget(subtitle)

        self.generate_btn = QPushButton("Сформировать КП")
        self.generate_btn.setObjectName("PrimaryButton")
        self.generate_btn.clicked.connect(self.generate)
        header.addLayout(header_text, stretch=1)
        header.addWidget(self.generate_btn, alignment=Qt.AlignTop)
        layout.addLayout(header)

        project_card = self.owner._card("Папка проекта")
        project_grid = QGridLayout()
        project_grid.setColumnStretch(1, 1)
        project_grid.setVerticalSpacing(12)
        project_grid.setHorizontalSpacing(10)
        project_card.layout().addLayout(project_grid)

        self.project_edit = QLineEdit()
        self.project_edit.setToolTip(self.project_dir_path)
        self.owner._add_row(
            project_grid,
            0,
            "Папка проекта",
            self.project_edit,
            "Выбрать",
            self.browse_project_dir,
        )
        layout.addWidget(project_card)

        files_card = self.owner._card("Файлы")
        files_grid = QGridLayout()
        files_grid.setColumnStretch(1, 1)
        files_grid.setVerticalSpacing(12)
        files_grid.setHorizontalSpacing(10)
        files_card.layout().addLayout(files_grid)

        self.calc_combo = QComboBox()
        self.calc_combo.setEditable(True)
        self.calc_combo.currentTextChanged.connect(self._on_calc_changed)

        self.sheet_combo = QComboBox()
        self.sheet_combo.setEditable(True)
        self.sheet_combo.currentTextChanged.connect(self._parse_excel)

        self.template_combo = QComboBox()
        self.template_combo.setEditable(True)

        self.output_edit = QLineEdit()
        self.output_edit.setToolTip(self.output_dir_path)

        self.owner._add_row(
            files_grid,
            0,
            "Excel-расчёт",
            self.calc_combo,
            "Обновить",
            lambda: self.scan_project(force=True),
        )
        self.owner._add_row(
            files_grid,
            1,
            "Лист Excel",
            self.sheet_combo,
            "Листы",
            self.load_sheets,
        )
        self.owner._add_row(
            files_grid,
            2,
            "Word-шаблон",
            self.template_combo,
            "Выбрать",
            self.browse_template_file,
        )
        self.owner._add_row(
            files_grid,
            3,
            "Папка результата",
            self.output_edit,
            "Выбрать",
            self.browse_output_dir,
        )
        layout.addWidget(files_card)

        data_card = self.owner._card("Данные КП")
        data_grid = QGridLayout()
        data_grid.setColumnStretch(1, 1)
        data_grid.setColumnStretch(3, 1)
        data_grid.setVerticalSpacing(12)
        data_grid.setHorizontalSpacing(10)
        data_card.layout().addLayout(data_grid)

        self.offer_version = QLineEdit()
        self.project_name = QLineEdit()
        self.delivery_terms = QLineEdit()
        self.delivery_time = QLineEdit()
        self.payment_terms = QLineEdit()

        data_grid.addWidget(QLabel("Версия:"), 0, 0)
        data_grid.addWidget(self.offer_version, 0, 1)
        data_grid.addWidget(QLabel("Проект:"), 0, 2)
        data_grid.addWidget(self.project_name, 0, 3)
        data_grid.addWidget(QLabel("Условия поставки:"), 1, 0)
        data_grid.addWidget(self.delivery_terms, 1, 1)
        data_grid.addWidget(QLabel("Срок поставки:"), 1, 2)
        data_grid.addWidget(self.delivery_time, 1, 3)
        data_grid.addWidget(QLabel("Условия оплаты:"), 2, 0)
        data_grid.addWidget(self.payment_terms, 2, 1, 1, 3)
        layout.addWidget(data_card)

        auto_card = self.owner._card("Автоматически заполняемые данные")
        self.auto_data_label = QLabel()
        self.auto_data_label.setWordWrap(True)
        self.auto_data_label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        auto_card.layout().addWidget(self.auto_data_label)
        layout.addWidget(auto_card)

        services_card = self.owner._card("Работы и услуги")
        services_layout = QHBoxLayout()
        self.engineering_check = QCheckBox("Инжиниринг включён")
        self.installation_check = QCheckBox("Монтаж включён")
        self.startup_check = QCheckBox("Пусконаладка включена")
        services_layout.addWidget(self.engineering_check)
        services_layout.addWidget(self.installation_check)
        services_layout.addWidget(self.startup_check)
        services_layout.addStretch(1)
        services_card.layout().addLayout(services_layout)
        layout.addWidget(services_card)

        items_card = self.owner._card("Позиции из Excel")
        self.items_table = QTableWidget(0, 5)
        self.items_table.setHorizontalHeaderLabels(
            ["В КП", "Наименование", "Кол-во", "Цена за ед.", "Сумма"]
        )
        self.items_table.verticalHeader().setVisible(False)
        self.items_table.setAlternatingRowColors(True)
        self.items_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.items_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        table_header = self.items_table.horizontalHeader()
        table_header.setSectionResizeMode(0, QHeaderView.ResizeToContents)
        table_header.setSectionResizeMode(1, QHeaderView.Stretch)
        for column in (2, 3, 4):
            table_header.setSectionResizeMode(column, QHeaderView.ResizeToContents)
        items_card.layout().addWidget(self.items_table)

        item_buttons = QHBoxLayout()
        parse_btn = QPushButton("Прочитать Excel")
        parse_btn.clicked.connect(self._parse_excel)
        all_btn = QPushButton("Выбрать все")
        all_btn.clicked.connect(lambda: self._set_all_items(True))
        none_btn = QPushButton("Снять все")
        none_btn.clicked.connect(lambda: self._set_all_items(False))
        item_buttons.addWidget(parse_btn)
        item_buttons.addWidget(all_btn)
        item_buttons.addWidget(none_btn)
        item_buttons.addStretch(1)
        items_card.layout().addLayout(item_buttons)
        layout.addWidget(items_card, stretch=1)

        self.status_label = QLabel("Выберите папку проекта")
        self.status_label.setWordWrap(True)
        self.status_label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        layout.addWidget(self.status_label)

        self.project_edit.textChanged.connect(self._on_project_text_changed)
        self.output_edit.textChanged.connect(self._on_output_text_changed)
        for widget in (
            self.offer_version,
            self.project_name,
            self.delivery_terms,
            self.delivery_time,
            self.payment_terms,
        ):
            widget.textChanged.connect(lambda *_: self.remember_values())
        self.template_combo.currentTextChanged.connect(lambda *_: self.remember_values())
        self.sheet_combo.currentTextChanged.connect(lambda *_: self.remember_values())

    # ------------------------------------------------------------ settings
    def _saved(self, key: str, default: str) -> str:
        value = self.settings.value(key, default)
        return str(value) if value is not None else default

    def _load_saved_values(self) -> None:
        self._set_line_path(self.project_edit, self.project_dir_path, is_file=False)
        self._set_line_path(self.output_edit, self.output_dir_path, is_file=False)

        self.offer_version.setText(self._saved("hvac/offer_version", "1"))
        self.project_name.setText(self._saved("hvac/project_name", ""))
        self.delivery_terms.setText(self._saved("hvac/delivery_terms", "DDP Алматы"))
        self.delivery_time.setText(self._saved("hvac/delivery_time", "16–20 недель"))
        self.payment_terms.setText(
            self._saved("hvac/payment_terms", "70% предоплата, 30% после поставки")
        )

        saved_calc = self._saved("hvac/calc_path", "")
        if saved_calc:
            self.owner._add_path_item(self.calc_combo, saved_calc, is_file=True)
        saved_sheet = self._saved("hvac/sheet_name", "")
        if saved_sheet:
            self.sheet_combo.addItem(saved_sheet)
            self.sheet_combo.setCurrentText(saved_sheet)
        saved_template = self._saved("hvac/template_path", "")
        if saved_template and Path(saved_template).is_file():
            self.owner._add_path_item(self.template_combo, saved_template, is_file=True)

    def remember_values(self) -> None:
        self.settings.setValue("hvac/project_dir", self.project_path_text())
        self.settings.setValue("hvac/output_dir", self.output_path_text())
        self.settings.setValue("hvac/calc_path", self._path_from_combo(self.calc_combo))
        self.settings.setValue("hvac/template_path", self._path_from_combo(self.template_combo))
        self.settings.setValue("hvac/sheet_name", self.sheet_combo.currentText().strip())
        self.settings.setValue("hvac/offer_version", self.offer_version.text().strip())
        self.settings.setValue("hvac/project_name", self.project_name.text().strip())
        self.settings.setValue("hvac/delivery_terms", self.delivery_terms.text().strip())
        self.settings.setValue("hvac/delivery_time", self.delivery_time.text().strip())
        self.settings.setValue("hvac/payment_terms", self.payment_terms.text().strip())
        self.settings.sync()

    def clear_cache(self) -> None:
        for key in (
            "hvac/project_dir",
            "hvac/output_dir",
            "hvac/calc_path",
            "hvac/template_path",
            "hvac/sheet_name",
            "hvac/project_name",
        ):
            self.settings.remove(key)
        self.project_dir_path = ""
        self.output_dir_path = ""
        self.client_company = ""
        self.from_client_dir = None
        self.basis_files = []
        self.parse_result = None
        self.items = []
        self.item_checks = []
        self._set_line_path(self.project_edit, "", is_file=False)
        self._set_line_path(self.output_edit, "", is_file=False)
        self.calc_combo.clear()
        self.sheet_combo.clear()
        self.items_table.setRowCount(0)
        self._load_default_template(clear_first=True)
        self._refresh_auto_data()
        self._update_status()

    def on_settings_changed(self) -> None:
        """Called by MainWindow after manager/signer settings are changed."""
        self._refresh_auto_data()

    # --------------------------------------------------------------- paths
    def _set_line_path(self, widget: QLineEdit, path_text: str, *, is_file: bool) -> None:
        self._updating_path_display = True
        try:
            self.owner._set_line_path(widget, path_text, is_file=is_file)
        finally:
            self._updating_path_display = False

    def _path_from_combo(self, combo: QComboBox) -> str:
        return self.owner._path_from_combo(combo)

    def project_path_text(self) -> str:
        return self.project_dir_path or self.project_edit.toolTip() or self.project_edit.text().strip()

    def output_path_text(self) -> str:
        return self.output_dir_path or self.output_edit.toolTip() or self.output_edit.text().strip()

    def _on_project_text_changed(self) -> None:
        if self._updating_path_display:
            return
        self.project_dir_path = self.project_edit.text().strip()
        self.project_edit.setToolTip(self.project_dir_path)
        self._refresh_project_metadata()
        self.remember_values()

    def _on_output_text_changed(self) -> None:
        if self._updating_path_display:
            return
        self.output_dir_path = self.output_edit.text().strip()
        self.output_edit.setToolTip(self.output_dir_path)
        self.remember_values()

    def browse_project_dir(self) -> None:
        start = self.project_path_text() or str(Path.home())
        path = QFileDialog.getExistingDirectory(self, "Выберите папку проекта", start)
        if not path:
            return
        self.project_dir_path = path
        self._set_line_path(self.project_edit, path, is_file=False)
        self.scan_project(force=True)

    def browse_output_dir(self) -> None:
        start = self.output_path_text() or self.project_path_text() or str(Path.home())
        path = QFileDialog.getExistingDirectory(self, "Выберите папку результата", start)
        if not path:
            return
        self.output_dir_path = path
        self._set_line_path(self.output_edit, path, is_file=False)
        self.remember_values()

    def browse_template_file(self) -> None:
        current = self._path_from_combo(self.template_combo)
        start = str(Path(current).parent) if current else self.project_path_text()
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите Word-шаблон HVAC",
            start,
            "Word (*.docx)",
        )
        if not path:
            return
        index = self.owner._find_combo_path(self.template_combo, path)
        if index < 0:
            self.owner._add_path_item(self.template_combo, path, is_file=True)
            index = self.template_combo.count() - 1
        self.template_combo.setCurrentIndex(index)
        self.remember_values()

    def _load_default_template(self, *, clear_first: bool = False) -> None:
        if clear_first:
            self.template_combo.clear()
        current = self._path_from_combo(self.template_combo)
        default_template = find_default_hvac_template()
        if default_template and self.owner._find_combo_path(self.template_combo, default_template) < 0:
            self.owner._add_path_item(self.template_combo, default_template, is_file=True)
        if current:
            current_index = self.owner._find_combo_path(self.template_combo, current)
            if current_index >= 0:
                self.template_combo.setCurrentIndex(current_index)
                return
        if default_template:
            default_index = self.owner._find_combo_path(self.template_combo, default_template)
            if default_index >= 0:
                self.template_combo.setCurrentIndex(default_index)

    # --------------------------------------------------------------- scan
    def scan_project(self, force: bool = False) -> None:
        project_text = self.project_path_text().strip()
        project_dir = Path(project_text) if project_text else None
        if not project_dir or not project_dir.is_dir():
            self.status_label.setText("Папка проекта не выбрана или недоступна")
            return

        if force:
            clear_scan_cache()
        found = scan_project_files(project_dir, use_cache=not force)
        excel_files = [Path(path) for path in found.get("excel", [])]
        old_calc = self._path_from_combo(self.calc_combo)

        self.calc_combo.blockSignals(True)
        try:
            self.calc_combo.clear()
            for path in excel_files:
                self.owner._add_path_item(self.calc_combo, str(path), is_file=True)

            selected_index = self.owner._find_combo_path(self.calc_combo, old_calc) if old_calc else -1
            if selected_index < 0 and excel_files:
                best = max(excel_files, key=_calc_file_score)
                selected_index = self.owner._find_combo_path(self.calc_combo, str(best))
            if selected_index >= 0:
                self.calc_combo.setCurrentIndex(selected_index)
        finally:
            self.calc_combo.blockSignals(False)

        if not self.output_path_text():
            self.output_dir_path = infer_output_dir(str(project_dir))
            self._set_line_path(self.output_edit, self.output_dir_path, is_file=False)

        self._refresh_project_metadata()
        self.load_sheets(refresh=False)
        self._parse_excel()
        self.remember_values()

    def _refresh_project_metadata(self) -> None:
        project_text = self.project_path_text().strip()
        if not project_text:
            self.client_company = ""
            self.from_client_dir = None
            self.basis_files = []
            self._refresh_auto_data()
            return

        self.client_company = extract_client_from_project_dir(project_text).strip()
        project_dir = Path(project_text)
        if project_dir.is_dir():
            self.from_client_dir = _find_from_client_dir(project_dir)
            self.basis_files = _list_basis_files(self.from_client_dir)
        else:
            self.from_client_dir = None
            self.basis_files = []
        self._refresh_auto_data()

    def _refresh_auto_data(self) -> None:
        signer = self._selected_signer()
        manager = self._manager_profile_dict()
        basis_folder = str(self.from_client_dir) if self.from_client_dir else "не найдена"
        basis_count = len(self.basis_files)
        self.auto_data_label.setText(
            "Дата: формируется автоматически\n"
            f"Заказчик: {self.client_company or 'не определён из пути'}\n"
            f"Подписант: {signer.get('name', '') or 'не выбран'} — "
            f"{signer.get('position', '')}\n"
            f"Исполнитель: {manager.get('name', '') or 'не заполнен в настройках'} — "
            f"{manager.get('position', '')}\n"
            f"Папка оснований: {basis_folder}\n"
            f"Документов для раздела «На основании»: {basis_count}"
        )

    # --------------------------------------------------------------- Excel
    def _on_calc_changed(self) -> None:
        self.load_sheets(refresh=False)
        self._parse_excel()
        self.remember_values()

    def load_sheets(self, refresh: bool = True) -> None:
        calc_path = self._path_from_combo(self.calc_combo)
        if not calc_path or not Path(calc_path).is_file():
            return
        old_sheet = self.sheet_combo.currentText().strip()
        try:
            names = read_sheet_names(calc_path)
        except Exception as exc:
            if refresh:
                QMessageBox.warning(self, "HVAC", f"Не удалось прочитать листы Excel:\n{exc}")
            return

        self.sheet_combo.blockSignals(True)
        try:
            self.sheet_combo.clear()
            self.sheet_combo.addItems(names)
            preferred = self._saved("hvac/sheet_name", "DDP_Almaty_20-10(v2)")
            for candidate in (old_sheet, preferred, "DDP_Almaty_20-10(v2)"):
                index = self.sheet_combo.findText(candidate)
                if index >= 0:
                    self.sheet_combo.setCurrentIndex(index)
                    break
        finally:
            self.sheet_combo.blockSignals(False)

    def _parse_excel(self) -> None:
        calc_path = self._path_from_combo(self.calc_combo)
        if not calc_path or not Path(calc_path).is_file():
            return
        try:
            result = parse_calculation(
                calc_path,
                sheet_name=self.sheet_combo.currentText().strip() or None,
                preferred_sheet="DDP_Almaty_20-10(v2)",
            )
        except Exception as exc:
            QMessageBox.warning(self, "HVAC", f"Не удалось разобрать calculation:\n{exc}")
            return

        self.parse_result = result
        self.items = list(result.items)
        self._apply_detected_delivery_terms(result.delivery_basis)
        if result.engineering.included is not None:
            self.engineering_check.setChecked(result.engineering.included)
        if result.installation.included is not None:
            self.installation_check.setChecked(result.installation.included)
        if result.startup.included is not None:
            self.startup_check.setChecked(result.startup.included)
        self._fill_items_table()
        self._update_status()
        self.remember_values()

    def _apply_detected_delivery_terms(self, detected: str | None) -> None:
        value = str(detected or "").strip()
        if not value:
            return
        current = self.delivery_terms.text().strip()
        if len(value) <= 4 and current.upper().startswith(value.upper()):
            return
        self.delivery_terms.setText(value)

    def _fill_items_table(self) -> None:
        self.items_table.setRowCount(0)
        self.item_checks.clear()
        for row, item in enumerate(self.items):
            self.items_table.insertRow(row)
            check = QCheckBox()
            check.setChecked(True)
            check.stateChanged.connect(self._update_status)
            self.item_checks.append(check)
            self.items_table.setCellWidget(row, 0, _centered_widget(check))
            self.items_table.setItem(row, 1, QTableWidgetItem(item.name))
            self.items_table.setItem(row, 2, QTableWidgetItem(format_qty(item.qty)))
            self.items_table.setItem(row, 3, _right_item(format_money(item.unit_price)))
            self.items_table.setItem(row, 4, _right_item(format_money(item.total_price)))
        self.items_table.resizeRowsToContents()

    def _set_all_items(self, checked: bool) -> None:
        for check in self.item_checks:
            check.setChecked(checked)
        self._update_status()

    def _selected_items(self) -> list[CalcItem]:
        return [item for item, check in zip(self.items, self.item_checks) if check.isChecked()]

    def _update_status(self) -> None:
        details: list[str] = []
        if self.client_company:
            details.append(f"Заказчик: {self.client_company}")
        if self.from_client_dir:
            details.append(f"оснований: {len(self.basis_files)}")
        else:
            details.append("папка From Client не найдена")

        if self.parse_result:
            selected = self._selected_items()
            total = sum(float(item.total_price or 0) for item in selected)
            details.extend(
                [
                    f"лист: {self.parse_result.sheet_name}",
                    f"позиций: {len(selected)} из {len(self.items)}",
                    f"итого: {format_money(total)} {self.parse_result.currency or ''}".strip(),
                ]
            )
            if self.parse_result.warnings:
                details.append("предупреждения: " + "; ".join(self.parse_result.warnings))
        else:
            details.append("calculation не прочитан")
        self.status_label.setText(" | ".join(details))

    # ------------------------------------------------------------- common data
    def _selected_signer(self) -> dict[str, str]:
        try:
            signer = self.owner._selected_signer()
            return {
                "name": str(signer.get("name", "")).strip(),
                "position": str(signer.get("position", "")).strip(),
            }
        except Exception:
            return {"name": "", "position": ""}

    def _manager_profile_dict(self) -> dict[str, str]:
        try:
            profile = self.owner._manager_profile()
            return {
                "name": str(getattr(profile, "name", "")).strip(),
                "position": str(getattr(profile, "position", "")).strip(),
                "email": str(getattr(profile, "email", "")).strip(),
                "phone": str(getattr(profile, "phone", "")).strip(),
            }
        except Exception:
            return {
                "name": self._saved("manager_name", ""),
                "position": self._saved("manager_position", ""),
                "email": self._saved("manager_email", ""),
                "phone": self._saved("manager_phone", ""),
            }

    # -------------------------------------------------------------- generate
    def generate(self) -> None:
        template_path = Path(self._path_from_combo(self.template_combo))
        if not template_path.is_file():
            QMessageBox.warning(self, "HVAC", "Не найден шаблон КП HVAC.")
            return

        project_path = self.project_path_text().strip()
        if not project_path or not Path(project_path).is_dir():
            QMessageBox.warning(self, "HVAC", "Сначала выберите папку проекта.")
            return

        self._refresh_project_metadata()
        if not self.client_company:
            QMessageBox.warning(
                self,
                "HVAC",
                "Не удалось определить заказчика из пути проекта. "
                "Проверьте, что путь содержит папку 02_Projects/КЛИЕНТ.",
            )
            return

        items = self._selected_items()
        if not items:
            QMessageBox.warning(self, "HVAC", "Не выбраны позиции для КП.")
            return

        output_dir = Path(self.output_path_text() or infer_output_dir(project_path))
        output_dir.mkdir(parents=True, exist_ok=True)
        version = _safe_filename(self.offer_version.text().strip() or "1")
        client = _safe_filename(self.client_company)
        output_path = output_dir / (
            f"Offer_{client}_{datetime.now():%d-%m-%y}(v{version}) HVAC.docx"
        )

        tags = self._collect_tags(items)
        try:
            _render_hvac_template(template_path, output_path, tags, items)
        except Exception as exc:
            QMessageBox.critical(self, "HVAC", f"Не удалось сформировать КП:\n{exc}")
            return

        self.remember_values()
        QMessageBox.information(self, "HVAC", f"КП сформировано:\n{output_path}")

    def _collect_tags(self, items: list[CalcItem]) -> dict[str, str]:
        result = self.parse_result
        currency = (result.currency if result else None) or "EUR"
        grand_total = sum(float(item.total_price or 0) for item in items)
        signer = self._selected_signer()
        manager = self._manager_profile_dict()

        return {
            "offer_date": datetime.now().strftime("%d.%m.%Y г."),
            "offer_version": self.offer_version.text().strip() or "1",
            "client_company_full": self.client_company,
            "intro_text": "",
            "project_name": self.project_name.text().strip(),
            "data_file_name": _format_basis_documents(self.basis_files),
            "delivery_terms": self.delivery_terms.text().strip(),
            "installation_terms": (
                "Монтажные работы включены."
                if self.installation_check.isChecked()
                else "Монтажные работы не включены."
            ),
            "startup_terms": (
                "Пуско-наладочные работы включены."
                if self.startup_check.isChecked()
                else "Пуско-наладочные работы не включены."
            ),
            "engineering_terms": (
                "Инжиниринг включен, срок выполнения от 7 недель."
                if self.engineering_check.isChecked()
                else "Инжиниринг не включен."
            ),
            "delivery_time": self.delivery_time.text().strip(),
            "unit_price_header": f"Цена за ед., {currency}",
            "total_price_header": f"Сумма, {currency}",
            "total_label": "Итого",
            "grand_total": format_money(grand_total),
            "total_price_block": f"{format_money(grand_total)} {currency}",
            "payment_terms": self.payment_terms.text().strip(),
            "signer_name": signer.get("name", ""),
            "signer_position": signer.get("position", ""),
            "manager_name": manager.get("name", ""),
            "manager_position": manager.get("position", ""),
            "manager_email": manager.get("email", ""),
            "manager_phone": manager.get("phone", ""),
        }


# ----------------------------------------------------------------- From Client
def _normalize_folder_name(name: str) -> str:
    text = str(name or "").replace("_", " ").replace("-", " ").casefold()
    text = re.sub(r"^\s*\d+[.)_\-\s]*", "", text)
    return " ".join(text.split())


def _is_from_client_name(name: str) -> bool:
    normalized = _normalize_folder_name(name)
    return any(alias in normalized for alias in _FROM_CLIENT_ALIASES)


def _find_from_client_dir(selected_dir: Path) -> Path | None:
    """Find 01_From Client near the selected project path.

    Users sometimes select the project root and sometimes a Sales docs folder,
    therefore direct children of the selected folder and its nearest parents are
    checked first. A shallow recursive fallback is used afterwards.
    """

    roots: list[Path] = []
    current = selected_dir
    for _ in range(3):
        if current not in roots and current.is_dir():
            roots.append(current)
        if current.parent == current:
            break
        current = current.parent

    direct_matches: list[Path] = []
    for root in roots:
        try:
            for child in root.iterdir():
                if child.is_dir() and _is_from_client_name(child.name):
                    direct_matches.append(child)
        except OSError:
            continue
    if direct_matches:
        return sorted(direct_matches, key=lambda path: (len(path.parts), path.name.casefold()))[0]

    # Keep network scans shallow: selected folder and its direct parent only.
    recursive_roots = roots[:2]
    recursive_matches: list[Path] = []
    for root in recursive_roots:
        try:
            for current_text, dir_names, _file_names in os.walk(root):
                current = Path(current_text)
                depth = len(current.parts) - len(root.parts)
                if depth >= 3:
                    dir_names[:] = []
                    continue
                for dir_name in list(dir_names):
                    candidate = current / dir_name
                    if _is_from_client_name(dir_name):
                        recursive_matches.append(candidate)
        except OSError:
            continue
    if not recursive_matches:
        return None
    return sorted(recursive_matches, key=lambda path: (len(path.parts), path.name.casefold()))[0]


def _list_basis_files(folder: Path | None) -> list[str]:
    if folder is None or not folder.is_dir():
        return []

    paths: list[Path] = []
    try:
        for path in folder.rglob("*"):
            try:
                if not path.is_file():
                    continue
            except OSError:
                continue
            if path.name.startswith("~$") or path.name.startswith("."):
                continue
            if path.name.casefold() in _IGNORED_BASIS_FILES:
                continue
            paths.append(path)
    except OSError:
        return []

    paths.sort(key=lambda path: str(path.relative_to(folder)).casefold())
    name_counts = Counter(path.name.casefold() for path in paths)
    result: list[str] = []
    for path in paths:
        if name_counts[path.name.casefold()] > 1:
            result.append(str(path.relative_to(folder)).replace("\\", "/"))
        else:
            result.append(path.name)
    return result


def _format_basis_documents(names: Iterable[str]) -> str:
    clean = [str(name).strip() for name in names if str(name).strip()]
    if not clean:
        return ""
    # The DOCX template already has a bullet before {{data_file_name}}.
    return clean[0] + "".join(f"\n• {name}" for name in clean[1:])


# -------------------------------------------------------------------- DOCX
def _render_hvac_template(
    template_path: Path,
    output_path: Path,
    tags: dict[str, str],
    items: list[CalcItem],
) -> None:
    document = Document(str(template_path))
    item_table = None
    template_row_index = None

    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            row_text = " | ".join(cell.text for cell in row.cells)
            found = {match.group(1).casefold() for match in _TAG_RE.finditer(row_text)}
            if found & _ITEM_TAGS:
                item_table = table
                template_row_index = row_index
                break
        if item_table is not None:
            break

    if item_table is None or template_row_index is None:
        raise ValueError("В шаблоне не найдена строка {{item_name}} ценовой таблицы.")

    template_row = item_table.rows[template_row_index]
    row_xml = deepcopy(template_row._tr)
    item_table._tbl.remove(template_row._tr)

    for offset, item in enumerate(items):
        insert_index = template_row_index + offset
        new_xml = deepcopy(row_xml)
        item_table._tbl.insert(insert_index, new_xml)
        new_row = item_table.rows[insert_index]
        item_tags = {
            "item_no": str(offset + 1),
            "item_name": item.name,
            "item_qty": format_qty(item.qty),
            "item_unit_price": format_money(item.unit_price),
            "item_total": format_money(item.total_price),
        }
        for cell in new_row.cells:
            for paragraph in cell.paragraphs:
                _replace_tags_in_paragraph(paragraph, item_tags)

    _replace_tags_everywhere(document, tags)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def _replace_tags_everywhere(document: Document, tags: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        _replace_tags_in_paragraph(paragraph, tags)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_tags_in_paragraph(paragraph, tags)
    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                _replace_tags_in_paragraph(paragraph, tags)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _replace_tags_in_paragraph(paragraph, tags)


def _replace_tags_in_paragraph(paragraph, tags: dict[str, str]) -> None:
    text = paragraph.text
    if "{{" not in text:
        return
    normalized_tags = {key.casefold(): str(value or "") for key, value in tags.items()}
    replaced = _TAG_RE.sub(
        lambda match: normalized_tags.get(match.group(1).casefold(), ""),
        text,
    )
    if replaced == text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replaced)


# ------------------------------------------------------------------ helpers
def _calc_file_score(path: Path) -> tuple[int, float, str]:
    name = path.name.casefold()
    score = 0
    if "calc" in name or "calculation" in name or "расчет" in name or "расчёт" in name:
        score += 30
    if "hvac" in name or "овкв" in name:
        score += 15
    if path.suffix.casefold() == ".xlsx":
        score += 2
    try:
        mtime = path.stat().st_mtime
    except OSError:
        mtime = 0.0
    return score, mtime, path.name.casefold()


def _right_item(text: str) -> QTableWidgetItem:
    item = QTableWidgetItem(text)
    item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
    return item


def _centered_widget(widget: QWidget) -> QWidget:
    box = QWidget()
    layout = QHBoxLayout(box)
    layout.setContentsMargins(0, 0, 0, 0)
    layout.addStretch(1)
    layout.addWidget(widget)
    layout.addStretch(1)
    return box


def _safe_filename(value: str) -> str:
    value = re.sub(r'[\\/:*?"<>|]+', "_", str(value or "")).strip()
    return "_".join(value.split()) or "HVAC"
