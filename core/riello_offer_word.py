from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
import re
from typing import Any

from docx import Document
from openpyxl import load_workbook

from core.docx_renderer import render_docx


@dataclass
class RielloOfferItem:
    model: str
    code: str = ""
    qty: float = 1.0
    unit_price: float = 0.0
    total: float = 0.0
    currency: str = "EUR"
    dimensions: str = ""
    weight_kg: float = 0.0
    note: str = ""
    kind: str = "Позиция"


@dataclass
class RielloOfferCalc:
    path: Path
    sheet_name: str
    currency: str = "EUR"
    items: list[RielloOfferItem] = field(default_factory=list)
    ddp_total: float | None = None
    ddp_label: str = ""
    warnings: list[str] = field(default_factory=list)

    @property
    def equipment(self) -> list[RielloOfferItem]:
        rows = [item for item in self.items if item.kind.lower().startswith("ибп")]
        return rows or list(self.items)

    @property
    def options(self) -> list[RielloOfferItem]:
        return [item for item in self.items if item.kind.lower().startswith("опц")]

    @property
    def source_total(self) -> float:
        return sum(float(item.total or 0.0) for item in self.items)


_REQUIRED_ITEM_TAGS = ("{{item_no}}", "{{item_name}}", "{{item_qty}}", "{{item_unit_price}}", "{{item_total}}")


def _number(value: Any, default: float = 0.0) -> float:
    if value is None or isinstance(value, bool):
        return default
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).replace("\u00a0", " ").replace(" ", "").replace(",", ".").strip()
    try:
        return float(text)
    except Exception:
        return default


def _fmt_qty(value: Any) -> str:
    number = _number(value, 0.0)
    return str(int(number)) if number.is_integer() else str(number).replace(".", ",")


def _fmt_money(value: Any) -> str:
    return f"{_number(value):,.2f}".replace(",", "TMP").replace(".", ",").replace("TMP", "\u00a0")


def _currency_name(currency: str) -> str:
    cur = str(currency or "").upper()
    return {"KZT": "тенге", "EUR": "евро", "USD": "долларов США"}.get(cur, cur or "")


def _offer_date() -> str:
    now = datetime.now()
    months = (
        "января", "февраля", "марта", "апреля", "мая", "июня",
        "июля", "августа", "сентября", "октября", "ноября", "декабря",
    )
    return f"{now.day} {months[now.month - 1]} {now.year} года"


def _sheet_starting(workbook, prefix: str):
    prefix = prefix.lower()
    for sheet in workbook.worksheets:
        if sheet.title.lower().startswith(prefix):
            return sheet
    return None


def _kind_from_note(note: str) -> str:
    lowered = str(note or "").lower()
    if "опция" in lowered:
        return "Опция"
    if "ибп" in lowered:
        return "ИБП"
    return "Позиция"


def _currency_from_header(value: Any) -> str:
    text = str(value or "").upper()
    match = re.search(r"\b(KZT|EUR|USD)\b", text)
    return match.group(1) if match else "EUR"


def _find_ddp_total(values_wb) -> tuple[float | None, str]:
    sheet = _sheet_starting(values_wb, "DDP")
    if sheet is None:
        return None, ""

    strongest: list[tuple[int, int, float, str]] = []
    for row in range(1, sheet.max_row + 1):
        labels = []
        for col in range(1, min(sheet.max_column, 4) + 1):
            value = sheet.cell(row, col).value
            if isinstance(value, str) and value.strip():
                labels.append(value.strip())
        label = " ".join(labels)
        normalized = label.lower().replace("ё", "е")
        numeric = None
        numeric_col = 0
        for col in range(sheet.max_column, 0, -1):
            value = sheet.cell(row, col).value
            if isinstance(value, (int, float)) and not isinstance(value, bool):
                numeric = float(value)
                numeric_col = col
                break
        if numeric is None or abs(numeric) < 1e-12:
            continue

        score = 0
        if any(token in normalized for token in ("final", "итого с ндс", "total with vat", "price with vat", "final price")):
            score += 100
        if any(token in normalized for token in ("итого", "total", "стоимость", "price")):
            score += 40
        if normalized.strip().startswith("ddp ") or " ddp " in f" {normalized} ":
            score += 30
        if score:
            strongest.append((score, row, numeric, label or f"DDP row {row}, col {numeric_col}"))

    if not strongest:
        return None, ""
    strongest.sort(key=lambda item: (item[0], item[1]), reverse=True)
    _score, _row, value, label = strongest[0]
    return value, label


def read_riello_offer_calc(path: str | Path) -> RielloOfferCalc:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Calc Riello не найден: {path}")

    formulas_wb = load_workbook(path, data_only=False, read_only=True)
    values_wb = load_workbook(path, data_only=True, read_only=True)
    try:
        sheet = _sheet_starting(formulas_wb, "UPS configuration")
        values_sheet = _sheet_starting(values_wb, "UPS configuration")
        if sheet is None or values_sheet is None:
            raise ValueError("В Calc не найден лист 'UPS configuration'. Выберите Calc, сформированный вкладкой Riello.")

        currency = _currency_from_header(sheet.cell(3, 8).value)
        result = RielloOfferCalc(path=path, sheet_name=sheet.title, currency=currency)

        for row in range(4, sheet.max_row + 1):
            model = str(sheet.cell(row, 2).value or "").strip()
            if not model:
                continue
            gross_marker = str(sheet.cell(row, 6).value or "").strip().lower()
            if model.upper() == "TOTAL" or gross_marker.startswith("gross total"):
                break

            code = str(sheet.cell(row, 3).value or "").strip()
            dimensions = str(sheet.cell(row, 4).value or "").strip()
            weight = _number(values_sheet.cell(row, 5).value, _number(sheet.cell(row, 5).value))
            unit_price = _number(values_sheet.cell(row, 6).value, _number(sheet.cell(row, 6).value))
            qty = _number(values_sheet.cell(row, 7).value, _number(sheet.cell(row, 7).value, 1.0)) or 1.0
            total = _number(values_sheet.cell(row, 8).value, 0.0)
            if total <= 0 and unit_price:
                total = unit_price * qty
            note = str(sheet.cell(row, 9).value or "").strip()
            result.items.append(
                RielloOfferItem(
                    model=model,
                    code=code,
                    qty=qty,
                    unit_price=unit_price,
                    total=total,
                    currency=currency,
                    dimensions=dimensions,
                    weight_kg=weight,
                    note=note,
                    kind=_kind_from_note(note),
                )
            )

        if not result.items:
            raise ValueError("На листе 'UPS configuration' не найдены позиции оборудования.")

        result.ddp_total, result.ddp_label = _find_ddp_total(values_wb)
        if result.ddp_total is None:
            result.warnings.append(
                "Итоговая DDP-стоимость не прочитана из кэша формул Excel. "
                "Если Calc только что создан, откройте его в Excel, сохраните и нажмите F5."
            )
        if len(result.equipment) > 1 and result.ddp_total is not None:
            result.warnings.append(
                "В Calc несколько ИБП. Общая DDP-сумма найдена, но Calc пока не содержит "
                "однозначных конечных цен по каждой позиции; в таблице КП используются цены строк UPS configuration."
            )
        return result
    finally:
        try:
            formulas_wb.close()
            values_wb.close()
        except Exception:
            pass


def preview_riello_offer_calc(path: str | Path) -> str:
    calc = read_riello_offer_calc(path)
    lines = [
        f"Calc: {calc.path.name}",
        f"Лист: {calc.sheet_name}",
        f"Валюта: {calc.currency}",
        f"Позиций: {len(calc.items)}",
        f"ИБП: {len(calc.equipment)}",
        f"Опций: {len(calc.options)}",
        "",
        "Позиции Calc:",
    ]
    for index, item in enumerate(calc.items, start=1):
        lines.append(
            f"{index}. [{item.kind}] {item.model} | код {item.code or '-'} | "
            f"кол-во {_fmt_qty(item.qty)} | цена {_fmt_money(item.unit_price)} {calc.currency} | "
            f"сумма {_fmt_money(item.total)} {calc.currency}"
        )
    lines.extend(["", f"Сумма строк оборудования: {_fmt_money(calc.source_total)} {calc.currency}"])
    if calc.ddp_total is not None:
        lines.append(f"Итог Calc/DDP: {_fmt_money(calc.ddp_total)} {calc.currency}")
        if calc.ddp_label:
            lines.append(f"Источник итога: {calc.ddp_label}")
    if calc.warnings:
        lines.extend(["", "Предупреждения:"])
        lines.extend(f"- {warning}" for warning in calc.warnings)
    return "\n".join(lines)


def validate_riello_word_template(path: str | Path) -> None:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError("Word-шаблон Riello не выбран или не найден.")
    document = Document(path)
    text_parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            text_parts.extend(cell.text for cell in row.cells)
    full_text = "\n".join(text_parts)
    missing = [tag for tag in _REQUIRED_ITEM_TAGS if tag not in full_text]
    if missing:
        raise ValueError(
            "Word-шаблон пока не размечен как шаблон КП. Не найдены теги таблицы: " + ", ".join(missing)
        )


def _offer_rows(calc: RielloOfferCalc) -> tuple[list[dict[str, Any]], float]:
    equipment = calc.equipment
    items: list[dict[str, Any]] = []

    # For the common Riello case one UPS configuration can safely receive the
    # final DDP total from Calc. Options remain part of its configuration.
    one_equipment_ddp = len(equipment) == 1 and calc.ddp_total is not None
    grand_total = calc.ddp_total if one_equipment_ddp else sum(item.total for item in equipment)

    for index, item in enumerate(equipment, start=1):
        total = calc.ddp_total if one_equipment_ddp else item.total
        unit_price = total / item.qty if item.qty else total
        name = item.model
        if item.code:
            name += f" ({item.code})"
        items.append(
            {
                "item_no": index,
                "item_name": name,
                "item_qty": _fmt_qty(item.qty),
                "item_unit_price": _fmt_money(unit_price),
                "item_total": _fmt_money(total),
            }
        )
    return items, float(grand_total or 0.0)


def generate_riello_word_offer(
    *,
    calc_path: str | Path,
    template_path: str | Path,
    output_dir: str | Path,
    client_name: str,
    signer_name: str = "",
    signer_position: str = "",
    manager_name: str = "",
    manager_position: str = "",
    manager_email: str = "",
    manager_phone: str = "",
) -> Path:
    validate_riello_word_template(template_path)
    calc = read_riello_offer_calc(calc_path)
    items, grand_total = _offer_rows(calc)
    if not items:
        raise ValueError("В Calc не найдены строки ИБП для ценовой таблицы КП.")

    output_dir = Path(output_dir)
    first_model = re.sub(r"[^A-Za-zА-Яа-я0-9._-]+", "_", calc.equipment[0].model).strip("_") or "Riello"
    safe_client = re.sub(r"[^A-Za-zА-Яа-я0-9._-]+", "_", client_name or "Client").strip("_") or "Client"
    output_path = output_dir / f"КП_Riello_{first_model}_{safe_client}_{datetime.now():%d-%m-%y}.docx"

    currency_name = _currency_name(calc.currency)
    options = [
        {
            "description": " — ".join(part for part in (item.model, item.code, item.note) if part),
            "qty": _fmt_qty(item.qty),
        }
        for item in calc.options
    ]
    technical_specs: list[dict[str, Any]] = []
    for item in calc.equipment:
        technical_specs.append({"name": f"ИБП {item.model}", "value": "", "is_section": True})
        if item.code:
            technical_specs.append({"name": "Код:", "value": item.code, "is_section": False})
        if item.dimensions:
            technical_specs.append({"name": "Габариты:", "value": item.dimensions, "is_section": False})
        if item.weight_kg:
            technical_specs.append({"name": "Вес:", "value": f"{_fmt_qty(item.weight_kg)} кг", "is_section": False})
        if item.note:
            technical_specs.append({"name": "Примечание:", "value": item.note, "is_section": False})

    replacements = {
        "{{offer_date}}": _offer_date(),
        "{{offer_version}}": "1",
        "{{client_company_full}}": client_name,
        "{{intro_text}}": "В ответ на Ваш запрос направляем коммерческое предложение на поставку оборудования Riello UPS.",
        "{{unit_price_header}}": f"Цена за единицу, {currency_name}",
        "{{total_price_header}}": f"Сумма, {currency_name}",
        "{{total_label}}": "ИТОГО",
        "{{grand_total}}": _fmt_money(grand_total),
        "{{total_price_block}}": f"{_fmt_money(grand_total)} {currency_name}.",
        "{{payment_terms}}": "Условия оплаты уточняются.",
        "{{delivery_time}}": "Срок поставки уточняется после размещения заказа.",
        "{{delivery_terms}}": "Согласно расчету Riello Calc.",
        "{{installation_terms}}": "Согласно расчету и составу предложения.",
        "{{startup_terms}}": "Согласно расчету и составу предложения.",
        "{{offer_validity}}": "Коммерческое предложение действительно в течение 30 календарных дней.",
        "{{currency_terms}}": f"Стоимость указана в {currency_name}.",
        "{{signer_name}}": signer_name,
        "{{signer_position}}": signer_position,
        "{{manager_name}}": manager_name,
        "{{manager_position}}": manager_position,
        "{{manager_email}}": manager_email,
        "{{manager_phone}}": manager_phone,
    }

    return render_docx(
        template_path=template_path,
        output_path=output_path,
        replacements=replacements,
        items=items,
        options=options,
        technical_specs=technical_specs,
    )
