from __future__ import annotations

from copy import deepcopy
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


BOLD_SPEC_ROWS = {
    "Тип модуля:",
    "Вентилятор:",
    "Компрессор:",
    "Выносной блок (Конденсор):",
}


def to_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def replace_in_paragraph(paragraph, replacements: dict[str, Any]) -> None:
    """Replace tags while preserving Word formatting where possible."""
    if not paragraph.runs:
        return

    changed = False
    for run in paragraph.runs:
        original = run.text
        updated = original
        for key, value in replacements.items():
            updated = updated.replace(key, to_text(value))
        if updated != original:
            run.text = updated
            changed = True

    if changed:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return

    updated_full = full_text
    for key, value in replacements.items():
        updated_full = updated_full.replace(key, to_text(value))

    if updated_full != full_text:
        paragraph.runs[0].text = updated_full
        for run in paragraph.runs[1:]:
            run.text = ""


def replace_tags(doc: Document, replacements: dict[str, Any]) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)


def set_cell_keep_style(cell, value: Any) -> None:
    value = to_text(value)
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
    else:
        paragraph = cell.add_paragraph()

    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)

    for extra_paragraph in cell.paragraphs[1:]:
        for run in extra_paragraph.runs:
            run.text = ""


def row_contains_tags(row, tags: list[str]) -> bool:
    row_text = "\n".join(cell.text for cell in row.cells)
    return any(tag in row_text for tag in tags)


def clone_row_after(table, source_row, after_row):
    new_tr = deepcopy(source_row._tr)
    after_row._tr.addnext(new_tr)
    return table.rows[after_row._index + 1]


def remove_row(table, row) -> None:
    table._tbl.remove(row._tr)


def fill_equipment_table(doc: Document, items: list[dict[str, Any]], total_label: str, grand_total: str) -> None:
    item_tags = ["{{item_no}}", "{{item_name}}", "{{item_qty}}", "{{item_unit_price}}", "{{item_total}}"]

    for table in doc.tables:
        template_row = None
        total_row = None
        for row in table.rows:
            if row_contains_tags(row, item_tags):
                template_row = row
            if row_contains_tags(row, ["{{total_label}}", "{{grand_total}}"]):
                total_row = row

        if template_row is None:
            continue

        insert_after = template_row
        for item in items:
            new_row = clone_row_after(table, template_row, insert_after)
            insert_after = new_row
            values = {
                "{{item_no}}": item.get("item_no", ""),
                "{{item_name}}": item.get("item_name", ""),
                "{{item_qty}}": item.get("item_qty", ""),
                "{{item_unit_price}}": item.get("item_unit_price", ""),
                "{{item_total}}": item.get("item_total", ""),
            }
            for cell in new_row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, values)
                remaining = cell.text
                for tag, value in values.items():
                    if tag in remaining:
                        set_cell_keep_style(cell, value)
                        break

        remove_row(table, template_row)

        if total_row is not None:
            values = {"{{total_label}}": total_label, "{{grand_total}}": grand_total}
            for cell in total_row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, values)
        return


def _paragraph_contains(paragraph, tag: str) -> bool:
    return tag in "".join(run.text for run in paragraph.runs) or tag in paragraph.text


def _clear_paragraph(paragraph) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = ""
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = ""


def _insert_table_after_paragraph(doc: Document, paragraph, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    paragraph._p.addnext(table._tbl)
    return table


def _find_tag_paragraph(doc: Document, tag: str):
    for paragraph in doc.paragraphs:
        if _paragraph_contains(paragraph, tag):
            return paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if _paragraph_contains(paragraph, tag):
                        return paragraph
    return None


def _set_cell_text(cell, text: Any, bold: bool = False) -> None:
    text = to_text(text)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for idx, part in enumerate(text.split("\n")):
        if idx:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        run = paragraph.add_run(part)
        run.bold = bold
        try:
            run.font.size = Pt(9)
        except Exception:
            pass


def _apply_table_style(table) -> None:
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.autofit = True


def insert_options_table(doc: Document, options: list[dict[str, Any]]) -> None:
    tag = "{{options_table}}"
    paragraph = _find_tag_paragraph(doc, tag)
    if not options:
        if paragraph is not None:
            _clear_paragraph(paragraph)
        return

    if paragraph is None:
        paragraph = doc.add_paragraph()
        paragraph.add_run(tag)

    table = _insert_table_after_paragraph(doc, paragraph, rows=1, cols=3)
    _apply_table_style(table)
    headers = ["№", "Наименование опции", "Кол-во"]
    for idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[idx], header, bold=True)

    for no, option in enumerate(options, start=1):
        row = table.add_row()
        _set_cell_text(row.cells[0], str(no))
        _set_cell_text(row.cells[1], option.get("description", ""))
        _set_cell_text(row.cells[2], option.get("qty", ""))

    _clear_paragraph(paragraph)


def insert_technical_specs_table(doc: Document, specs: list[dict[str, Any]]) -> None:
    tag = "{{technical_specs_table}}"
    paragraph = _find_tag_paragraph(doc, tag)
    if not specs:
        if paragraph is not None:
            _clear_paragraph(paragraph)
        return

    if paragraph is None:
        paragraph = doc.add_paragraph()
        paragraph.add_run(tag)

    table = _insert_table_after_paragraph(doc, paragraph, rows=0, cols=2)
    _apply_table_style(table)

    for spec in specs:
        row = table.add_row()
        if spec.get("is_section"):
            merged = row.cells[0].merge(row.cells[1])
            _set_cell_text(merged, spec.get("name", ""), bold=True)
        else:
            _set_cell_text(row.cells[0], spec.get("name", ""), bold=True)
            _set_cell_text(row.cells[1], spec.get("value", ""))

    _clear_paragraph(paragraph)



# --- Template-row based STULZ specification rendering -----------------------

def _element_parent(element):
    return element.getparent()


def _remove_element(element) -> None:
    parent = _element_parent(element)
    if parent is not None:
        parent.remove(element)


def _insert_element_after(after_element, new_element):
    after_element.addnext(new_element)
    return new_element


def _make_empty_paragraph_element():
    return OxmlElement("w:p")


def _make_page_break_paragraph_element():
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "page")
    run.append(br)
    paragraph.append(run)
    return paragraph


def _find_table_with_tags(doc: Document, tags: list[str]):
    for table in doc.tables:
        for row in table.rows:
            if row_contains_tags(row, tags):
                return table
    return None


def _replace_tags_in_table(table, replacements: dict[str, Any]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)
            for tag, value in replacements.items():
                if tag in cell.text:
                    set_cell_keep_style(cell, value)


def set_option_name_cell_text(cell, text: str) -> None:
    """
    Форматирование ячейки опции:

    первая строка:
    - жирный
    - 11 pt

    последующий текст:
    - НЕ жирный
    - 10 pt
    - переносы сохраняются через Word line break
    """
    text = to_text(text).strip()
    lines = [line.strip() for line in text.splitlines()]

    title = lines[0] if lines else ""
    body_lines = [line for line in lines[1:] if line]

    # Полностью очищаем ячейку, чтобы не наследовать жирный run из шаблона.
    cell.text = ""

    paragraph = cell.paragraphs[0]

    # Дополнительно сбрасываем форматирование всех старых run, если Word их сохранил.
    for run in paragraph.runs:
        run.text = ""
        run.bold = False
        run.font.bold = False
        run.font.size = Pt(10)

    if title:
        title_run = paragraph.add_run(title)
        title_run.bold = True
        title_run.font.bold = True
        title_run.font.size = Pt(11)

    if body_lines:
        body_run = paragraph.add_run()
        body_run.bold = False
        body_run.font.bold = False
        body_run.font.size = Pt(10)

        for line in body_lines:
            body_run.add_break(WD_BREAK.LINE)
            body_run.add_text(line)

        # Важно: python-docx иногда наследует bold из шаблона.
        # Поэтому после добавления текста ещё раз проходим по run и принудительно
        # оставляем жирным только первый run с заголовком.
        for i, run in enumerate(paragraph.runs):
            if i == 0 and title:
                run.bold = True
                run.font.bold = True
                run.font.size = Pt(11)
            else:
                run.bold = False
                run.font.bold = False
                run.font.size = Pt(10)


def _insert_pdf_first_page_as_picture(paragraph, pdf_path: Any, width_inches: float = 6.5) -> bool:
    """Render first PDF page to PNG and insert it into the given paragraph.

    Requires PyMuPDF package: pip install PyMuPDF
    """
    if not pdf_path:
        return False

    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        return False

    try:
        import fitz  # PyMuPDF
    except Exception:
        paragraph.add_run(f"[PDF-чертёж не вставлен: не установлен PyMuPDF]")
        return False

    tmp_png_path = None
    try:
        with fitz.open(str(pdf_path)) as pdf_doc:
            if pdf_doc.page_count < 1:
                return False

            page = pdf_doc.load_page(0)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp_png_path = Path(tmp.name)
            pix.save(str(tmp_png_path))

        run = paragraph.add_run()
        run.add_picture(str(tmp_png_path), width=Inches(width_inches))
        return True
    except Exception as exc:
        paragraph.add_run(f"[PDF-чертёж не вставлен: {exc}]")
        return False
    finally:
        if tmp_png_path and tmp_png_path.exists():
            try:
                tmp_png_path.unlink()
            except Exception:
                pass


def _fill_template_row_table(table, row_tags: list[str], rows: list[dict[str, Any]]) -> None:
    template_row = None
    for row in table.rows:
        if row_contains_tags(row, row_tags):
            template_row = row
            break
    if template_row is None:
        return

    insert_after = template_row
    for values in rows:
        new_row = clone_row_after(table, template_row, insert_after)
        insert_after = new_row
        replacements = {tag: values.get(tag, "") for tag in row_tags}
        for cell in new_row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)
            for tag, value in replacements.items():
                if tag in cell.text:
                    if tag == "{{opt_name}}":
                        set_option_name_cell_text(cell, value)
                    else:
                        set_cell_keep_style(cell, value)
                    break
    remove_row(table, template_row)


def _fill_options_template_table(table, options: list[dict[str, Any]]) -> None:
    rows = []
    for idx, option in enumerate(options, start=1):
        rows.append({
            "{{opt_no}}": str(idx),
            "{{opt_name}}": option.get("description", ""),
            "{{opt_qty}}": option.get("qty", ""),
        })
    _fill_template_row_table(table, ["{{opt_no}}", "{{opt_name}}", "{{opt_qty}}"], rows)


def _fill_specs_template_table(table, specs: list[dict[str, Any]]) -> None:
    template_row = None
    for row in table.rows:
        if row_contains_tags(row, ["{{technical_specs_parameter}}", "{{technical_specs_value}}"]):
            template_row = row
            break
    if template_row is None:
        return

    insert_after = template_row
    for spec in specs:
        new_row = clone_row_after(table, template_row, insert_after)
        insert_after = new_row

        parameter = spec.get("name", "")
        value = "" if spec.get("is_section") else spec.get("value", "")
        values = {
            "{{technical_specs_parameter}}": parameter,
            "{{technical_specs_value}}": value,
        }
        is_bold = parameter.strip() in BOLD_SPEC_ROWS

        for cell in new_row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, values)
                for run in paragraph.runs:
                    run.bold = is_bold
            for tag, replacement_value in values.items():
                if tag in cell.text:
                    set_cell_keep_style(cell, replacement_value)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = is_bold
                    break

    remove_row(table, template_row)


def insert_stulz_specification_blocks(doc: Document, spec_blocks: list[dict[str, Any]]) -> bool:
    """Fill Word template-row specification blocks for one or many STULZ models.

    Order for each model:
    options title -> options table -> one empty line ->
    technical specs title -> technical specs table -> optional drawing PDF image.

    The drawing PDF is taken from block["drawing_pdf"] and inserted after
    the technical specs table.
    """
    option_title_p = _find_tag_paragraph(doc, "{{options_title}}")
    option_table = _find_table_with_tags(doc, ["{{opt_no}}", "{{opt_name}}", "{{opt_qty}}"])
    specs_title_p = _find_tag_paragraph(doc, "{{technical_specs_title}}")
    specs_table = _find_table_with_tags(doc, ["{{technical_specs_parameter}}", "{{technical_specs_value}}"])

    if option_title_p is None or option_table is None or specs_title_p is None or specs_table is None:
        return False

    anchor = specs_table._tbl

    from docx.text.paragraph import Paragraph
    from docx.table import Table

    for block_index, block in enumerate(spec_blocks):
        option_title_el = deepcopy(option_title_p._p)
        option_table_el = deepcopy(option_table._tbl)
        specs_title_el = deepcopy(specs_title_p._p)
        specs_table_el = deepcopy(specs_table._tbl)

        drawing_pdf = block.get("drawing_pdf")
        drawing_paragraph_el = _make_empty_paragraph_element() if drawing_pdf else None

        elements = [
            option_title_el,
            option_table_el,
            _make_empty_paragraph_element(),
            specs_title_el,
            specs_table_el,
        ]

        if drawing_paragraph_el is not None:
            elements.append(_make_empty_paragraph_element())
            elements.append(drawing_paragraph_el)

        # One empty paragraph between different conditioner blocks.
        if block_index < len(spec_blocks) - 1:
            elements.append(_make_empty_paragraph_element())

        for el in elements:
            anchor = _insert_element_after(anchor, el)

        option_title_clone = Paragraph(option_title_el, option_title_p._parent)
        option_table_clone = Table(option_table_el, option_table._parent)
        specs_title_clone = Paragraph(specs_title_el, specs_title_p._parent)
        specs_table_clone = Table(specs_table_el, specs_table._parent)

        replace_in_paragraph(option_title_clone, {"{{options_title}}": block.get("options_title", "")})
        _fill_options_template_table(option_table_clone, block.get("options", []))
        replace_in_paragraph(specs_title_clone, {"{{technical_specs_title}}": block.get("technical_specs_title", "")})
        _fill_specs_template_table(specs_table_clone, block.get("technical_specs", []))

        if drawing_paragraph_el is not None:
            drawing_paragraph = Paragraph(drawing_paragraph_el, specs_title_p._parent)
            _insert_pdf_first_page_as_picture(drawing_paragraph, drawing_pdf)

    # Remove the original template block after cloned blocks have been inserted.
    for element in (option_title_p._p, option_table._tbl, specs_title_p._p, specs_table._tbl):
        _remove_element(element)
    return True

def remove_empty_service_tags(doc: Document) -> None:
    replace_tags(doc, {
        "{{options_table}}": "",
        "{{technical_specs_table}}": "",
        "{{options_title}}": "",
        "{{opt_no}}": "",
        "{{opt_name}}": "",
        "{{opt_qty}}": "",
        "{{technical_specs_title}}": "",
        "{{technical_specs_parameter}}": "",
        "{{technical_specs_value}}": "",
    })


def get_unique_output_path(path: Path) -> Path:
    if not path.exists():
        return path
    parent = path.parent
    stem = path.stem
    suffix = path.suffix
    counter = 1
    while True:
        candidate = parent / f"{stem}_{counter}{suffix}"
        if not candidate.exists():
            return candidate
        counter += 1


def render_docx(
    template_path: str | Path,
    output_path: str | Path,
    replacements: dict[str, Any],
    items: list[dict[str, Any]] | None = None,
    options: list[dict[str, Any]] | None = None,
    technical_specs: list[dict[str, Any]] | None = None,
    stulz_spec_blocks: list[dict[str, Any]] | None = None,
) -> Path:
    template_path = Path(template_path)
    output_path = Path(output_path)
    doc = Document(template_path)

    fill_equipment_table(
        doc=doc,
        items=items or [],
        total_label=to_text(replacements.get("{{total_label}}", "ИТОГО")),
        grand_total=to_text(replacements.get("{{grand_total}}", "")),
    )

    replace_tags(doc, replacements)
    used_template_blocks = False
    if stulz_spec_blocks:
        used_template_blocks = insert_stulz_specification_blocks(doc, stulz_spec_blocks)
    if not used_template_blocks:
        insert_options_table(doc, options or [])
        insert_technical_specs_table(doc, technical_specs or [])
    remove_empty_service_tags(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path = get_unique_output_path(output_path)
    doc.save(output_path)
    return output_path
