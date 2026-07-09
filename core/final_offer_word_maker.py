from __future__ import annotations

"""Universal DOCX commercial-offer builder.

The module deliberately contains no brand, Excel or GUI logic.  A caller gives
it a Word template, ordinary tag values and zero or more repeating table rows.
The same builder can therefore be used by HVAC, Stulz, Eltek, Riello and future
brands as their templates are migrated to common {{tag_name}} placeholders.
"""

from copy import deepcopy
from dataclasses import dataclass, field
import os
from pathlib import Path
import re
import tempfile
from typing import Any, Iterable, Mapping, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Row
from docx.text.paragraph import Paragraph


_TAG_RE = re.compile(r"\{\{\s*([A-Za-zА-Яа-я0-9_\-.]+)\s*\}\}")


@dataclass(frozen=True)
class RepeatingTable:
    """Description of one repeatable Word table row.

    ``row_marker`` is any tag contained in the template row, for example
    ``item_name`` or ``{{item_name}}``.  Each item in ``rows`` is a mapping of
    tag names to values.  The template row is cloned for every item and then
    removed.
    """

    row_marker: str
    rows: Sequence[Mapping[str, Any]]
    required: bool = True
    keep_rows_together: bool = True
    repeat_header: bool = True


@dataclass
class OfferBuildResult:
    output_path: Path
    replaced_tags: set[str] = field(default_factory=set)
    unresolved_tags: set[str] = field(default_factory=set)
    created_table_rows: int = 0
    warnings: list[str] = field(default_factory=list)




_CURRENCY_WORDS = {
    "KZT": (("тенге", "тенге", "тенге"), ("тиын", "тиына", "тиынов")),
    "EUR": (("евро", "евро", "евро"), ("евроцент", "евроцента", "евроцентов")),
    "USD": (("доллар США", "доллара США", "долларов США"), ("цент", "цента", "центов")),
    "RUB": (("рубль", "рубля", "рублей"), ("копейка", "копейки", "копеек")),
}


def format_amount_in_words(value: Any, currency: str = "KZT") -> str:
    """Return a formatted amount with a full Russian text spelling.

    Example: ``92 932,88 EUR (Девяносто две тысячи девятьсот тридцать два
    евро 88 евроцентов)``. The function has no brand-specific dependencies and
    can be reused by every offer page.
    """

    try:
        amount = round(float(value), 2)
    except (TypeError, ValueError) as exc:
        raise ValueError(f"Некорректная сумма: {value!r}") from exc
    negative = amount < 0
    amount = abs(amount)
    major = int(amount)
    minor = int(round((amount - major) * 100))
    if minor == 100:
        major += 1
        minor = 0
    code = str(currency or "KZT").upper()
    major_forms, minor_forms = _CURRENCY_WORDS.get(code, ((code, code, code), ("", "", "")))
    words = _integer_to_russian_words(major)
    if negative:
        words = "минус " + words
    major_word = _plural_form(major, major_forms)
    minor_word = _plural_form(minor, minor_forms).strip()
    numeric = f"{('-' if negative else '')}{major:,}.{minor:02d}".replace(",", " ").replace(".", ",")
    spelled = f"{words} {major_word} {minor:02d} {minor_word}".strip()
    return f"{numeric} {code} ({spelled[:1].upper() + spelled[1:]})"


def _plural_form(number: int, forms: Sequence[str]) -> str:
    value = abs(int(number)) % 100
    if 11 <= value <= 19:
        return forms[2]
    value %= 10
    if value == 1:
        return forms[0]
    if 2 <= value <= 4:
        return forms[1]
    return forms[2]


def _integer_to_russian_words(number: int) -> str:
    if number == 0:
        return "ноль"
    units_m = ("", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять")
    units_f = ("", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять")
    teens = ("десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать")
    tens = ("", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят", "семьдесят", "восемьдесят", "девяносто")
    hundreds = ("", "сто", "двести", "триста", "четыреста", "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот")
    scales = (
        ("", "", "", False),
        ("тысяча", "тысячи", "тысяч", True),
        ("миллион", "миллиона", "миллионов", False),
        ("миллиард", "миллиарда", "миллиардов", False),
        ("триллион", "триллиона", "триллионов", False),
    )
    if number >= 10 ** (3 * len(scales)):
        return str(number)
    parts: list[str] = []
    group_index = 0
    while number:
        group = number % 1000
        number //= 1000
        if group:
            chunk: list[str] = []
            chunk.append(hundreds[group // 100])
            rest = group % 100
            if 10 <= rest <= 19:
                chunk.append(teens[rest - 10])
            else:
                chunk.append(tens[rest // 10])
                units = units_f if scales[group_index][3] else units_m
                chunk.append(units[rest % 10])
            if group_index:
                chunk.append(_plural_form(group, scales[group_index][:3]))
            parts.insert(0, " ".join(word for word in chunk if word))
        group_index += 1
    return " ".join(parts)


class OfferTemplateError(ValueError):
    """Raised when a required structural marker is absent from a template."""


def make_final_offer(
    template_path: str | Path,
    output_path: str | Path,
    tags: Mapping[str, Any] | None = None,
    repeating_tables: Sequence[RepeatingTable] | None = None,
    *,
    clear_unresolved: bool = True,
    atomic_save: bool = True,
    overwrite: bool = False,
) -> OfferBuildResult:
    """Fill a DOCX template and save the final commercial offer.

    Tag keys may be passed either as ``offer_date`` or ``{{offer_date}}``.
    Replacements are case-insensitive.  Placeholders split by Word into several
    runs are replaced without flattening the whole paragraph, which preserves
    the formatting of neighbouring text.
    """

    template = Path(template_path)
    output = Path(output_path)
    if not template.is_file():
        raise FileNotFoundError(f"Шаблон Word не найден: {template}")
    if template.suffix.casefold() != ".docx":
        raise OfferTemplateError(f"Ожидается DOCX-шаблон: {template}")
    if output.suffix.casefold() != ".docx":
        output = output.with_suffix(".docx")
    if output.exists() and not overwrite:
        raise FileExistsError(f"Файл уже существует: {output}")

    document = Document(str(template))
    result = OfferBuildResult(output_path=output)
    scalar_tags, paragraph_list_tags = _split_tag_values(tags or {})

    for block in repeating_tables or ():
        created = _fill_repeating_table(document, block, result)
        result.created_table_rows += created

    for tag_name, items in paragraph_list_tags.items():
        created = _replace_tag_with_paragraph_list(document, tag_name, items)
        if created:
            result.replaced_tags.add(tag_name)
        else:
            result.warnings.append(
                f"В Word-шаблоне не найден отдельный абзац с тегом {{{{{tag_name}}}}}."
            )

    replaced = _replace_mapping_everywhere(document, scalar_tags)
    result.replaced_tags.update(replaced)

    unresolved = collect_unresolved_tags(document)
    result.unresolved_tags = set(unresolved)
    if unresolved:
        result.warnings.append(
            "В шаблоне остались незаполненные теги: "
            + ", ".join(sorted(f"{{{{{name}}}}}" for name in unresolved))
        )
        if clear_unresolved:
            empty_values = {name.casefold(): "" for name in unresolved}
            result.replaced_tags.update(_replace_mapping_everywhere(document, empty_values))

    output.parent.mkdir(parents=True, exist_ok=True)
    if atomic_save:
        _save_atomic(document, output)
    else:
        document.save(str(output))
    _verify_saved_file(output)
    return result


def collect_unresolved_tags(document: DocumentObject) -> set[str]:
    """Return placeholder names still present anywhere in the DOCX."""

    result: set[str] = set()
    for paragraph in _iter_all_paragraphs(document):
        for match in _TAG_RE.finditer(_paragraph_text(paragraph)):
            result.add(match.group(1).strip())
    return result


def normalize_tag_name(value: Any) -> str:
    text = str(value or "").strip()
    match = _TAG_RE.fullmatch(text)
    if match:
        text = match.group(1)
    return text.strip().casefold()


def _normalize_mapping(values: Mapping[str, Any]) -> dict[str, str]:
    normalized: dict[str, str] = {}
    for key, value in values.items():
        name = normalize_tag_name(key)
        if name:
            normalized[name] = "" if value is None else str(value)
    return normalized


def _split_tag_values(
    values: Mapping[str, Any],
) -> tuple[dict[str, str], dict[str, list[str]]]:
    """Split scalar tags from values that must become repeated paragraphs."""

    scalar: dict[str, str] = {}
    paragraph_lists: dict[str, list[str]] = {}
    for key, value in values.items():
        name = normalize_tag_name(key)
        if not name:
            continue
        if isinstance(value, (list, tuple, set)):
            paragraph_lists[name] = [
                str(item).strip() for item in value if str(item).strip()
            ]
        else:
            scalar[name] = "" if value is None else str(value)
    return scalar, paragraph_lists


def _replace_tag_with_paragraph_list(
    document: DocumentObject,
    tag_name: str,
    items: Sequence[str],
) -> int:
    """Clone a standalone tag paragraph once per list item.

    The cloned paragraphs keep the bullet/numbering, indentation, spacing and
    text formatting defined in the Word template.
    """

    marker = normalize_tag_name(tag_name)
    if not marker:
        return 0

    matches = []
    for paragraph in list(_iter_all_paragraphs(document)):
        text = _paragraph_text(paragraph).strip()
        found = _TAG_RE.fullmatch(text)
        if found and found.group(1).strip().casefold() == marker:
            matches.append(paragraph)

    created_total = 0
    for paragraph in matches:
        if not items:
            _replace_mapping_in_paragraph(paragraph, {marker: ""})
            continue

        template_xml = deepcopy(paragraph._p)
        _replace_mapping_in_paragraph(paragraph, {marker: items[0]})
        created_total += 1
        insert_after = paragraph._p

        for item in items[1:]:
            new_xml = deepcopy(template_xml)
            insert_after.addnext(new_xml)
            insert_after = new_xml
            new_paragraph = Paragraph(new_xml, paragraph._parent)
            _replace_mapping_in_paragraph(new_paragraph, {marker: item})
            created_total += 1

    return created_total


def _fill_repeating_table(
    document: DocumentObject,
    block: RepeatingTable,
    result: OfferBuildResult,
) -> int:
    marker = normalize_tag_name(block.row_marker)
    if not marker:
        raise OfferTemplateError("Для повторяемой таблицы не указан row_marker.")

    found = _find_template_row(document, marker)
    if found is None:
        message = f"В Word-шаблоне не найдена строка с тегом {{{{{marker}}}}}."
        if block.required:
            raise OfferTemplateError(message)
        result.warnings.append(message)
        return 0

    table, template_row, template_row_index = found
    if block.repeat_header and template_row_index > 0:
        _set_repeat_table_header(table.rows[template_row_index - 1]._tr)
    source_xml = deepcopy(template_row._tr)
    insert_after = template_row._tr
    created = 0

    for raw_values in block.rows:
        row_values = _normalize_mapping(raw_values)
        new_xml = deepcopy(source_xml)
        if block.keep_rows_together:
            _set_row_cant_split(new_xml)
        insert_after.addnext(new_xml)
        insert_after = new_xml
        new_row = _Row(new_xml, table)
        result.replaced_tags.update(_replace_mapping_in_row(new_row, row_values))
        created += 1

    table._tbl.remove(template_row._tr)
    return created


def _find_template_row(document: DocumentObject, marker: str):
    for table in _iter_all_tables(document):
        for row_index, row in enumerate(table.rows):
            row_text = "\n".join(cell.text for cell in row.cells)
            names = {match.group(1).strip().casefold() for match in _TAG_RE.finditer(row_text)}
            if marker in names:
                return table, row, row_index
    return None



def _set_row_cant_split(row_xml) -> None:
    tr_pr = row_xml.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_repeat_table_header(row_xml) -> None:
    tr_pr = row_xml.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))

def _replace_mapping_in_row(row, values: Mapping[str, str]) -> set[str]:
    replaced: set[str] = set()
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            replaced.update(_replace_mapping_in_paragraph(paragraph, values))
        for nested_table in cell.tables:
            for nested_row in nested_table.rows:
                replaced.update(_replace_mapping_in_row(nested_row, values))
    return replaced


def _replace_mapping_everywhere(
    document: DocumentObject,
    values: Mapping[str, str],
) -> set[str]:
    replaced: set[str] = set()
    if not values:
        return replaced
    for paragraph in _iter_all_paragraphs(document):
        replaced.update(_replace_mapping_in_paragraph(paragraph, values))
    return replaced


def _replace_mapping_in_paragraph(paragraph, values: Mapping[str, str]) -> set[str]:
    full_text = _paragraph_text(paragraph)
    if "{{" not in full_text:
        return set()

    matches = list(_TAG_RE.finditer(full_text))
    replaced: set[str] = set()
    # Replace from right to left so positions of earlier placeholders stay valid.
    for match in reversed(matches):
        raw_name = match.group(1).strip()
        key = raw_name.casefold()
        if key not in values:
            continue
        if _replace_range_across_runs(paragraph, match.start(), match.end(), values[key]):
            replaced.add(raw_name)
    return replaced


def _replace_range_across_runs(paragraph, start: int, end: int, replacement: str) -> bool:
    if start < 0 or end < start:
        return False
    if not paragraph.runs:
        text = paragraph.text
        paragraph.text = text[:start] + replacement + text[end:]
        return True

    spans: list[tuple[int, int, Any]] = []
    position = 0
    for run in paragraph.runs:
        run_start = position
        run_end = position + len(run.text)
        spans.append((run_start, run_end, run))
        position = run_end

    involved = [item for item in spans if item[1] > start and item[0] < end]
    if not involved:
        return False

    first_start, _first_end, first_run = involved[0]
    last_start, _last_end, last_run = involved[-1]
    prefix = first_run.text[: max(0, start - first_start)]
    suffix = last_run.text[max(0, end - last_start) :]

    if first_run is last_run:
        first_run.text = prefix + replacement + suffix
    else:
        first_run.text = prefix + replacement
        for _run_start, _run_end, run in involved[1:-1]:
            run.text = ""
        last_run.text = suffix
    return True


def _paragraph_text(paragraph) -> str:
    if paragraph.runs:
        return "".join(run.text for run in paragraph.runs)
    return paragraph.text or ""


def _iter_all_tables(document: DocumentObject):
    seen: set[Any] = set()

    def emit(tables):
        for table in tables:
            identity = table._tbl
            if identity in seen:
                continue
            seen.add(identity)
            yield table
            for row in table.rows:
                for cell in row.cells:
                    yield from emit(cell.tables)

    yield from emit(document.tables)
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from emit(container.tables)


def _iter_all_paragraphs(document: DocumentObject):
    seen: set[Any] = set()

    def emit_paragraphs(paragraphs):
        for paragraph in paragraphs:
            identity = paragraph._p
            if identity not in seen:
                seen.add(identity)
                yield paragraph

    def emit_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from emit_paragraphs(cell.paragraphs)
                    yield from emit_tables(cell.tables)

    yield from emit_paragraphs(document.paragraphs)
    yield from emit_tables(document.tables)
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from emit_paragraphs(container.paragraphs)
            yield from emit_tables(container.tables)


def _save_atomic(document: DocumentObject, output: Path) -> None:
    temp_path: Path | None = None
    try:
        fd, raw_temp_path = tempfile.mkstemp(
            suffix=".docx",
            prefix=f".{output.stem}_",
            dir=str(output.parent),
        )
        os.close(fd)
        temp_path = Path(raw_temp_path)
        document.save(str(temp_path))
        _verify_saved_file(temp_path)
        os.replace(temp_path, output)
    finally:
        if temp_path is not None and temp_path.exists():
            try:
                temp_path.unlink()
            except OSError:
                pass


def _verify_saved_file(path: Path) -> None:
    try:
        stat = path.stat()
    except OSError as exc:
        raise OSError(f"Не удалось подтвердить сохранение Word-файла: {path}") from exc
    if not path.is_file() or stat.st_size <= 0:
        raise OSError(f"Word-файл не создан или имеет нулевой размер: {path}")


__all__ = [
    "OfferBuildResult",
    "OfferTemplateError",
    "RepeatingTable",
    "collect_unresolved_tags",
    "make_final_offer",
    "normalize_tag_name",
]
