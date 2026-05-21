from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


@dataclass
class ManagerProfile:
    name: str = ""
    position: str = ""
    email: str = ""
    phone: str = ""

    def is_empty(self) -> bool:
        return not any([self.name.strip(), self.position.strip(), self.email.strip(), self.phone.strip()])


EMAIL_RE = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE)
PHONE_RE = re.compile(r"(?:\+?\d[\d\s()\-]{7,}\d)")
POSITION_WORDS = (
    "менеджер",
    "инженер",
    "специалист",
    "директор",
    "руководитель",
    "начальник",
    "консультант",
    "координатор",
)


def normalize_line(text: str) -> str:
    return " ".join(str(text or "").replace("\xa0", " ").split()).strip(" :;\t")


def docx_lines(path: Path) -> list[str]:
    if Document is None:
        return []
    try:
        doc = Document(str(path))
    except Exception:
        return []

    lines: list[str] = []
    for paragraph in doc.paragraphs:
        line = normalize_line(paragraph.text)
        if line:
            lines.append(line)

    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_line(cell.text) for cell in row.cells]
            for cell in cells:
                if cell:
                    lines.append(cell)
            joined = normalize_line(" ".join(cells))
            if joined:
                lines.append(joined)

    return lines


def candidate_docx_files(project_dir: Path) -> list[Path]:
    if not project_dir.exists():
        return []

    files = [p for p in project_dir.rglob("*.docx") if p.is_file() and not p.name.startswith("~$")]

    def score(path: Path) -> tuple[int, float]:
        name = path.name.lower()
        value = 0
        if "кп" in name or "offer" in name or "офер" in name or "commercial" in name:
            value += 20
        if "template" in name or "шаблон" in name or "tagged" in name:
            value -= 10
        try:
            mtime = path.stat().st_mtime
        except OSError:
            mtime = 0.0
        return (value, mtime)

    return sorted(files, key=score, reverse=True)[:20]


def extract_manager_from_lines(lines: Iterable[str]) -> ManagerProfile:
    clean = [normalize_line(line) for line in lines]
    clean = [line for line in clean if line]
    lower = [line.lower() for line in clean]

    start = -1
    for i, line in enumerate(lower):
        if "исполнитель" in line:
            start = i
            break
    if start < 0:
        return ManagerProfile()

    # Usually the block is immediately after the word "Исполнитель".
    block = clean[start : min(len(clean), start + 16)]
    if block and block[0].lower().startswith("исполнитель"):
        after_label = normalize_line(re.sub("исполнитель", "", block[0], flags=re.IGNORECASE))
        block = ([after_label] if after_label else []) + block[1:]

    profile = ManagerProfile()

    for line in block:
        if not profile.email:
            m = EMAIL_RE.search(line)
            if m:
                profile.email = m.group(0)

        if not profile.phone:
            m = PHONE_RE.search(line)
            if m:
                profile.phone = normalize_line(m.group(0))

    for line in block:
        low = line.lower()
        if profile.email and profile.email.lower() in low:
            continue
        if profile.phone and profile.phone in line:
            continue
        if "исполнитель" in low:
            continue
        if not profile.position and any(word in low for word in POSITION_WORDS):
            profile.position = line
            continue

    for line in block:
        low = line.lower()
        if "исполнитель" in low:
            continue
        if EMAIL_RE.search(line) or PHONE_RE.search(line):
            continue
        if profile.position and line == profile.position:
            continue
        words = line.split()
        # A Russian full name is usually 2-4 capitalized words. Keep this permissive.
        if not profile.name and 2 <= len(words) <= 5 and not any(ch.isdigit() for ch in line):
            profile.name = line
            break

    return profile


def parse_manager_from_docx(path: Path) -> ManagerProfile:
    return extract_manager_from_lines(docx_lines(path))


def find_manager_in_project(project_dir: Path) -> ManagerProfile:
    for path in candidate_docx_files(project_dir):
        profile = parse_manager_from_docx(path)
        if not profile.is_empty():
            return profile
    return ManagerProfile()
