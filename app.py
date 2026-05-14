from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import load_workbook

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

MONTHS_RU = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая", 6: "июня",
    7: "июля", 8: "августа", 9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}

@dataclass
class CalcData:
    sheet_name: str
    version: str
    model: str
    quantity: float
    unit_price: float
    total_price: float
    delivery_basis: str
    options: list[tuple[str, float]]


def money(value: float) -> str:
    return f"{value:,.2f}".replace(",", " ").replace(".", ",")


def sanitize_filename(value: str) -> str:
    value = re.sub(r"[\\/:*?\"<>|]+", "_", value).strip()
    return value or "client"


def load_config(project_dir: Path) -> dict:
    cfg = project_dir / "config.json"
    if not cfg.exists():
        cfg = project_dir / "config.example.json"
    if cfg.exists():
        return json.loads(cfg.read_text(encoding="utf-8"))
    return {}


def first_not_empty(*values):
    for value in values:
        if value not in (None, ""):
            return value
    return ""


def as_float(value, default: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return default
        return float(value)
    except Exception:
        return default


def find_row_by_label(ws, label: str) -> Optional[int]:
    label_norm = label.lower().strip()
    for row in ws.iter_rows():
        for cell in row:
            val = cell.value
            if isinstance(val, str) and val.lower().strip() == label_norm:
                return cell.row
    return None


def parse_calc(xlsx_path: Path, sheet_name: Optional[str] = None) -> CalcData:
    wb_values = load_workbook(xlsx_path, data_only=True)
    wb_formulas = load_workbook(xlsx_path, data_only=False)
    ws = wb_values[sheet_name] if sheet_name else wb_values[wb_values.sheetnames[0]]
    wsf = wb_formulas[ws.title]

    version = str(first_not_empty(ws["C1"].value, "Version 1"))
    model = str(first_not_empty(ws["D2"].value, ws["E2"].value, "Equipment"))
    quantity = as_float(ws["C3"].value, 1.0)

    total_row = find_row_by_label(ws, "TOTAL")
    total_price = as_float(ws.cell(total_row, 4).value if total_row else None, 0.0)
    unit_price = total_price / quantity if quantity else total_price

    delivery_basis = "DDP г. Алматы" if "DDP" in ws.title.upper() else "EXW Hamburg, Germany"

    options: list[tuple[str, float]] = []
    start = find_row_by_label(ws, "Options:") or 6
    end = find_row_by_label(ws, "Total EXW for us") or min(start + 80, ws.max_row)
    for row_idx in range(start + 1, end):
        name = ws.cell(row_idx, 1).value
        qty = ws.cell(row_idx, 3).value
        if isinstance(name, str) and name.strip() and as_float(qty, 0) > 0:
            options.append((name.strip(), as_float(qty)))

    # fallback: if Excel was not calculated, use visible zero but keep model/qty
    if total_price == 0:
        # Try to take price from row named DDP/EXW if it has a cached value.
        for label in ("DDP Almaty", "EXW - Hamburg, Germany", "Total per quantity"):
            r = find_row_by_label(ws, label)
            if r and as_float(ws.cell(r, 4).value) > 0:
                total_price = as_float(ws.cell(r, 4).value)
                unit_price = total_price / quantity if quantity else total_price
                break

    return CalcData(
        sheet_name=ws.title,
        version=version,
        model=model,
        quantity=quantity,
        unit_price=unit_price,
        total_price=total_price,
        delivery_basis=delivery_basis,
        options=options,
    )


def iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    text = paragraph.text
    new_text = text
    for key, value in replacements.items():
        new_text = new_text.replace(key, value)
    if new_text != text:
        # Preserve simple paragraph style; rebuild runs for reliable replacement across split runs.
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


def replace_text(document: Document, replacements: dict[str, str]) -> None:
    for p in iter_paragraphs(document):
        replace_in_paragraph(p, replacements)


def find_table_by_headers(document: Document, required_headers: Iterable[str]):
    headers = [h.lower() for h in required_headers]
    for table in document.tables:
        table_text = " | ".join(cell.text.lower() for row in table.rows[:2] for cell in row.cells)
        if all(h in table_text for h in headers):
            return table
    return None


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def fill_product_table(document: Document, calc: CalcData, currency: str) -> None:
    table = find_table_by_headers(document, ["наименование", "кол-во", "цена"])
    if not table:
        return
    # Prefer first data row after headers. Original template has one blank row.
    row = table.rows[1] if len(table.rows) > 1 else table.add_row()
    while len(row.cells) < 5:
        table.add_row()
        row = table.rows[-1]
    set_cell(row.cells[0], "1")
    set_cell(row.cells[1], calc.model)
    set_cell(row.cells[2], str(int(calc.quantity) if calc.quantity.is_integer() else calc.quantity))
    set_cell(row.cells[3], money(calc.unit_price))
    set_cell(row.cells[4], money(calc.total_price))

    # Try to update total row if present.
    for r in table.rows:
        if "итого" in " ".join(c.text.lower() for c in r.cells):
            set_cell(r.cells[-1], money(calc.total_price), bold=True)
            break


def fill_spec_table(document: Document, calc: CalcData) -> None:
    table = find_table_by_headers(document, ["наименование опции", "кол-во"])
    if not table or not calc.options:
        return
    # Keep header row, remove most existing body content by blanking rows, then repopulate.
    while len(table.rows) > 1:
        tbl = table._tbl
        tbl.remove(table.rows[-1]._tr)
    for idx, (name, qty) in enumerate(calc.options, start=1):
        row = table.add_row()
        set_cell(row.cells[0], str(idx))
        set_cell(row.cells[1], name)
        qty_text = str(int(qty) if float(qty).is_integer() else qty)
        set_cell(row.cells[2], f"{qty_text} шт.")


def extract_pdf_text(pdf_path: Path, max_chars: int = 5000) -> str:
    if not pdf_path.exists() or fitz is None:
        return ""
    chunks = []
    with fitz.open(pdf_path) as doc:
        for page in doc:
            chunks.append(page.get_text("text"))
            if sum(len(x) for x in chunks) > max_chars:
                break
    return "\n".join(chunks).strip()[:max_chars]


def append_pdf_description(document: Document, pdf_dir: Optional[Path], model: str) -> None:
    if not pdf_dir or not pdf_dir.exists():
        return
    candidates = [pdf_dir / f"{model}.pdf"] + list(pdf_dir.glob(f"*{model}*.pdf"))
    pdf_path = next((p for p in candidates if p.exists()), None)
    if not pdf_path:
        return
    text = extract_pdf_text(pdf_path)
    if not text:
        return
    document.add_page_break()
    document.add_heading("Техническое описание", level=1)
    document.add_paragraph(text)


def make_offer(template_path: Path, calc_path: Path, output_dir: Path, client_name: str,
               sheet_name: Optional[str] = None, pdf_dir: Optional[Path] = None,
               version: Optional[str] = None, city: str = "г. Алматы") -> Path:
    calc = parse_calc(calc_path, sheet_name)
    document = Document(template_path)
    now = datetime.now()
    date_ru = f"{city}, {now.day} {MONTHS_RU[now.month]} {now.year} г."
    version_value = version or calc.version.replace("Version", "Версия №")

    total_words = ""
    replacements = {
        "ТОО «[Организация]»": f"ТОО «{client_name}»",
        "[Организация]": client_name,
        "г. Алматы, 15 апреля 2021 г.": date_ru,
        "Версия №1": version_value,
        "в течение 13-14 недель после поступления предоплаты.": "в течение 13-14 недель после поступления предоплаты.",
        "DDP г. Алматы": calc.delivery_basis,
        "Итого, стоимость оборудования составляет:  ( ) EUR 00 eurocent, с учетом НДС 12%.":
            f"Итого, стоимость оборудования составляет: {money(calc.total_price)} EUR, с учетом НДС 12%.",
        "кондиционеров CCU121A": f"кондиционеров {calc.model}",
        "Тип модуля:\n\tASU 211 AL": f"Тип модуля:\n\t{calc.model}",
    }
    replace_text(document, replacements)
    fill_product_table(document, calc, "EUR")
    fill_spec_table(document, calc)
    append_pdf_description(document, pdf_dir, calc.model)

    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"КП_{sanitize_filename(client_name)}_{calc.model}_{now:%Y-%m-%d}.docx"
    out = output_dir / filename
    document.save(out)
    return out


def run_gui(project_dir: Path) -> None:
    try:
        from PySide6.QtCore import Qt, QSettings
        from PySide6.QtGui import QFont, QIcon
        from PySide6.QtWidgets import (
            QApplication,
            QComboBox,
            QFileDialog,
            QFrame,
            QGridLayout,
            QHBoxLayout,
            QLabel,
            QLineEdit,
            QMainWindow,
            QMessageBox,
            QPushButton,
            QSizePolicy,
            QSpacerItem,
            QTextEdit,
            QVBoxLayout,
            QWidget,
        )
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(
            "Для запуска красивого GUI установите PySide6: pip install PySide6"
        ) from exc

    cfg = load_config(project_dir)

    class SamOfferWindow(QMainWindow):
        def __init__(self) -> None:
            super().__init__()
            self.settings = QSettings("SAM Group", "SAM Offer Generator")
            self.setWindowTitle("SAM Offer Generator")
            self.setMinimumSize(1040, 680)
            self.setWindowIcon(QIcon())

            self.client_edit = QLineEdit(self._saved("client", "ТОО Example"))
            self.template_edit = QLineEdit(self._saved_path("template", cfg.get("default_template", "templates/kp_template.docx")))
            self.calc_edit = QLineEdit(self._saved_path("calc", cfg.get("default_calc", "samples/Calc_23-12-24 PAC.xlsx")))
            self.pdf_edit = QLineEdit(self._saved_path("pdf_dir", cfg.get("default_pdf_dir", "pdf")))
            self.output_edit = QLineEdit(self._saved_path("output_dir", cfg.get("default_output_dir", "output")))
            self.sheet_combo = QComboBox()
            self.sheet_combo.setEditable(True)
            self.status_label = QLabel("Готов к формированию коммерческого предложения")
            self.preview = QTextEdit()
            self.preview.setReadOnly(True)
            self.preview.setPlaceholderText("Здесь появится краткая проверка Excel-файла и результат генерации.")

            self._build_ui()
            self._apply_sam_style()
            self._load_sheets()
            self._refresh_preview()

        def _saved(self, key: str, default: str) -> str:
            value = self.settings.value(key, default)
            return str(value) if value is not None else default

        def _saved_path(self, key: str, default_relative: str) -> str:
            default_path = str(project_dir / default_relative)
            value = self.settings.value(key, default_path)
            return str(value) if value is not None else default_path

        def _build_ui(self) -> None:
            central = QWidget()
            root = QHBoxLayout(central)
            root.setContentsMargins(0, 0, 0, 0)
            root.setSpacing(0)

            sidebar = QFrame()
            sidebar.setObjectName("Sidebar")
            sidebar.setFixedWidth(290)
            side = QVBoxLayout(sidebar)
            side.setContentsMargins(28, 34, 28, 28)
            side.setSpacing(18)

            brand = QLabel("SAM\nGROUP")
            brand.setObjectName("Brand")
            brand.setAlignment(Qt.AlignLeft)
            title = QLabel("Offer Generator")
            title.setObjectName("SideTitle")
            subtitle = QLabel("Коммерческие предложения из Excel, Word и PDF")
            subtitle.setObjectName("SideSubtitle")
            subtitle.setWordWrap(True)

            badge = QLabel("MVP → Windows EXE")
            badge.setObjectName("Badge")
            badge.setAlignment(Qt.AlignCenter)

            side.addWidget(brand)
            side.addWidget(title)
            side.addWidget(subtitle)
            side.addSpacing(12)
            side.addWidget(badge)
            side.addSpacerItem(QSpacerItem(20, 20, QSizePolicy.Minimum, QSizePolicy.Expanding))
            side.addWidget(QLabel("SAM Group style\nDark graphite · Signal red · Clean white"))

            content = QFrame()
            content.setObjectName("Content")
            content_layout = QVBoxLayout(content)
            content_layout.setContentsMargins(34, 28, 34, 28)
            content_layout.setSpacing(18)

            header = QHBoxLayout()
            h_text = QVBoxLayout()
            page_title = QLabel("Новое коммерческое предложение")
            page_title.setObjectName("PageTitle")
            page_subtitle = QLabel("Заполните клиента, выберите расчет и шаблон. Остальное программа соберет автоматически.")
            page_subtitle.setObjectName("PageSubtitle")
            h_text.addWidget(page_title)
            h_text.addWidget(page_subtitle)
            self.generate_btn = QPushButton("Сформировать КП")
            self.generate_btn.setObjectName("PrimaryButton")
            self.generate_btn.clicked.connect(self._generate)
            header.addLayout(h_text)
            header.addWidget(self.generate_btn, alignment=Qt.AlignTop)
            content_layout.addLayout(header)

            form_card = self._card("Данные предложения")
            grid = QGridLayout()
            form_card.layout().addLayout(grid)
            grid.setColumnStretch(1, 1)
            grid.setVerticalSpacing(12)
            grid.setHorizontalSpacing(10)

            rows = [
                ("Клиент", self.client_edit, None),
                ("Word-шаблон", self.template_edit, lambda: self._browse_file(self.template_edit, "Word (*.docx)")),
                ("Excel-расчет", self.calc_edit, lambda: self._browse_file(self.calc_edit, "Excel (*.xlsx)")),
                ("Папка PDF", self.pdf_edit, lambda: self._browse_dir(self.pdf_edit)),
                ("Папка результата", self.output_edit, lambda: self._browse_dir(self.output_edit)),
                ("Лист Excel", self.sheet_combo, None),
            ]
            for row, (label, widget, command) in enumerate(rows, start=1):
                lab = QLabel(label)
                lab.setObjectName("FormLabel")
                grid.addWidget(lab, row, 0)
                grid.addWidget(widget, row, 1)
                if command:
                    btn = QPushButton("Выбрать")
                    btn.setObjectName("GhostButton")
                    btn.clicked.connect(command)
                    grid.addWidget(btn, row, 2)
                elif widget is self.sheet_combo:
                    btn = QPushButton("Обновить")
                    btn.setObjectName("GhostButton")
                    btn.clicked.connect(self._load_sheets)
                    grid.addWidget(btn, row, 2)

            content_layout.addWidget(form_card)

            bottom = QHBoxLayout()
            preview_card = self._card("Проверка данных")
            preview_card.layout().addWidget(self.preview)
            bottom.addWidget(preview_card, stretch=2)

            status_card = self._card("Статус")
            status_text = QLabel(
                "1. Выберите Excel-калькуляцию\n"
                "2. Проверьте лист и клиента\n"
                "3. Нажмите кнопку формирования\n\n"
                "Результат сохраняется в output/ как .docx"
            )
            status_text.setWordWrap(True)
            status_card.layout().addWidget(status_text)
            status_card.layout().addWidget(self.status_label)
            bottom.addWidget(status_card, stretch=1)
            content_layout.addLayout(bottom)

            root.addWidget(sidebar)
            root.addWidget(content)
            self.setCentralWidget(central)

            for widget in (self.client_edit, self.template_edit, self.calc_edit, self.pdf_edit, self.output_edit):
                widget.textChanged.connect(self._refresh_preview)
            self.sheet_combo.currentTextChanged.connect(self._refresh_preview)
            self.calc_edit.textChanged.connect(self._load_sheets)

        def _card(self, title: str) -> QFrame:
            frame = QFrame()
            frame.setObjectName("Card")
            layout = QVBoxLayout(frame)
            layout.setContentsMargins(20, 18, 20, 20)
            layout.setSpacing(12)
            label = QLabel(title)
            label.setObjectName("CardTitle")
            layout.addWidget(label)
            return frame

        def _apply_sam_style(self) -> None:
            app = QApplication.instance()
            if app:
                app.setFont(QFont("Segoe UI", 10))
            self.setStyleSheet("""
                QMainWindow { background: #F4F6F8; }
                #Sidebar { background: #15171B; color: #FFFFFF; }
                #Content { background: #F4F6F8; }
                #Brand { color: #FFFFFF; font-size: 34px; font-weight: 900; letter-spacing: 2px; }
                #SideTitle { color: #FFFFFF; font-size: 23px; font-weight: 700; }
                #SideSubtitle { color: #B8C0CC; font-size: 13px; line-height: 1.4; }
                #Badge { background: #D71920; color: white; border-radius: 16px; padding: 8px 12px; font-weight: 700; }
                #PageTitle { color: #171A1F; font-size: 28px; font-weight: 800; }
                #PageSubtitle { color: #667085; font-size: 13px; }
                #Card { background: #FFFFFF; border: 1px solid #E7EAF0; border-radius: 18px; }
                #CardTitle { color: #171A1F; font-size: 15px; font-weight: 800; }
                #FormLabel { color: #344054; font-weight: 700; }
                QLineEdit, QComboBox, QTextEdit {
                    background: #FFFFFF; border: 1px solid #D0D5DD; border-radius: 10px;
                    padding: 9px 11px; color: #101828; selection-background-color: #D71920;
                }
                QLineEdit:focus, QComboBox:focus, QTextEdit:focus { border: 1px solid #D71920; }
                QTextEdit { min-height: 150px; }
                QPushButton { border-radius: 11px; padding: 10px 16px; font-weight: 800; }
                #PrimaryButton { background: #D71920; color: white; border: 1px solid #D71920; }
                #PrimaryButton:hover { background: #B9151B; }
                #GhostButton { background: #FFFFFF; color: #1D2939; border: 1px solid #D0D5DD; }
                #GhostButton:hover { border: 1px solid #D71920; color: #D71920; }
                QLabel { color: #475467; }
            """)

        def _browse_file(self, target: QLineEdit, file_filter: str) -> None:
            path, _ = QFileDialog.getOpenFileName(self, "Выберите файл", target.text(), file_filter)
            if path:
                target.setText(path)
                if target is self.calc_edit:
                    self._load_sheets()

        def _browse_dir(self, target: QLineEdit) -> None:
            path = QFileDialog.getExistingDirectory(self, "Выберите папку", target.text())
            if path:
                target.setText(path)

        def _load_sheets(self) -> None:
            current = self.sheet_combo.currentText().strip()
            self.sheet_combo.blockSignals(True)
            self.sheet_combo.clear()
            try:
                calc_path = Path(self.calc_edit.text())
                if calc_path.exists():
                    wb = load_workbook(calc_path, read_only=True, data_only=True)
                    self.sheet_combo.addItems(wb.sheetnames)
                    if current and current in wb.sheetnames:
                        self.sheet_combo.setCurrentText(current)
            except Exception:
                self.sheet_combo.addItem("")
            finally:
                self.sheet_combo.blockSignals(False)

        def _refresh_preview(self) -> None:
            try:
                calc_path = Path(self.calc_edit.text())
                if not calc_path.exists():
                    self.preview.setPlainText("Excel-файл пока не найден.")
                    return
                calc = parse_calc(calc_path, self.sheet_combo.currentText().strip() or None)
                qty = int(calc.quantity) if calc.quantity.is_integer() else calc.quantity
                lines = [
                    f"Лист: {calc.sheet_name}",
                    f"Версия: {calc.version}",
                    f"Модель: {calc.model}",
                    f"Количество: {qty}",
                    f"Условия поставки: {calc.delivery_basis}",
                    f"Итого: {money(calc.total_price)} EUR",
                    f"Опций найдено: {len(calc.options)}",
                ]
                self.preview.setPlainText("\n".join(lines))
            except Exception as exc:
                self.preview.setPlainText(f"Не удалось прочитать расчет: {exc}")

        def _remember_values(self) -> None:
            self.settings.setValue("client", self.client_edit.text())
            self.settings.setValue("template", self.template_edit.text())
            self.settings.setValue("calc", self.calc_edit.text())
            self.settings.setValue("pdf_dir", self.pdf_edit.text())
            self.settings.setValue("output_dir", self.output_edit.text())

        def _generate(self) -> None:
            try:
                self.generate_btn.setEnabled(False)
                self.status_label.setText("Формирую документ...")
                QApplication.processEvents()
                out = make_offer(
                    template_path=Path(self.template_edit.text()),
                    calc_path=Path(self.calc_edit.text()),
                    output_dir=Path(self.output_edit.text()),
                    client_name=self.client_edit.text().strip() or "Client",
                    sheet_name=self.sheet_combo.currentText().strip() or None,
                    pdf_dir=Path(self.pdf_edit.text()) if self.pdf_edit.text().strip() else None,
                )
                self._remember_values()
                self.status_label.setText(f"Готово: {out.name}")
                QMessageBox.information(self, "SAM Offer Generator", f"КП сформировано:\n{out}")
            except Exception as exc:
                self.status_label.setText("Ошибка формирования")
                QMessageBox.critical(self, "Ошибка", str(exc))
            finally:
                self.generate_btn.setEnabled(True)
                self._refresh_preview()

    app = QApplication.instance() or QApplication(sys.argv)
    window = SamOfferWindow()
    window.show()
    app.exec()

def main(argv=None) -> int:
    project_dir = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
    parser = argparse.ArgumentParser(description="Генератор коммерческого предложения в Word")
    parser.add_argument("--gui", action="store_true", help="Запустить окно Windows")
    parser.add_argument("--template", default=str(project_dir / "templates/kp_template.docx"))
    parser.add_argument("--calc", default=str(project_dir / "samples/Calc_23-12-24 PAC.xlsx"))
    parser.add_argument("--pdf-dir", default=str(project_dir / "pdf"))
    parser.add_argument("--output-dir", default=str(project_dir / "output"))
    parser.add_argument("--client", default="ТОО Example")
    parser.add_argument("--sheet", default=None)
    args = parser.parse_args(argv)

    if args.gui or len(sys.argv) == 1:
        run_gui(project_dir)
        return 0

    out = make_offer(
        template_path=Path(args.template),
        calc_path=Path(args.calc),
        output_dir=Path(args.output_dir),
        client_name=args.client,
        sheet_name=args.sheet,
        pdf_dir=Path(args.pdf_dir) if args.pdf_dir else None,
    )
    print(out)
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
